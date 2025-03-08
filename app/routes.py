# app/routes.py

from flask import render_template, flash, redirect, request, session, url_for, send_file, send_from_directory, request, g, jsonify, Blueprint, Flask
from flask_login import login_required, current_user, login_user, logout_user
from functools import wraps
from .extensions import supabase, logger, stripe  # Removed oauth import
from .celery_app import celery  # Import the Celery instance
from .models import User
#from .generate_query import generate_job_keywords, generate_urls
from .jobmatcher import embed_user_preferences, calculate_user_job_fit
from forms import JobPreferencesForm, EducationEntryForm
from math import ceil
from .tasks import process_job_preferences
from datetime import datetime, timedelta, timezone
from decouple import config
import json
import ast
import os
import subprocess
import requests
#from urllib.parse import quote_plus, urlencode
from dateutil import parser

import time
import openai
from openai import AzureOpenAI
import base64
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re
from lxml import etree

main_bp = Blueprint('main', __name__)
app = Flask(__name__)

client_id = os.getenv("AZURE_CLIENT_ID")
tenant_id = os.getenv("AZURE_TENANT_ID")
api_key = os.getenv("AZURE_OPENAI_API_KEY_COVER_LETTER")
client_secret = os.getenv("AZURE_CLIENT_SECRET")
azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT_COVER_LETTER")
openai.api_key = os.getenv("OPENAI_API_KEY")


# Check for missing variables
if not all([api_key, azure_endpoint]):
    raise EnvironmentError("Ensure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT_COVER_LETTER are set.")

# Initialize AzureOpenAI client with API Key
client = AzureOpenAI(
    api_key=api_key,
    azure_endpoint=azure_endpoint,
    api_version="2024-05-01-preview"
)

def time_ago(created_at):
    # Parse the created_at string into a datetime object
    created_at = datetime.strptime(created_at, '%Y-%m-%dT%H:%M:%S.%f%z')

    # Get the current time as a timezone-aware datetime
    now = datetime.now(timezone.utc)
    

    # Calculate the time difference
    delta = now - created_at

    # Convert time delta to seconds
    seconds = delta.total_seconds()

    if seconds < 60:
        return f"{int(seconds)} seconds ago"
    elif seconds < 3600:  # less than an hour
        minutes = int(seconds // 60)
        return f"{minutes} minute{'s' if minutes > 1 else ''} ago"
    elif seconds < 86400:  # less than a day
        hours = int(seconds // 3600)
        return f"{hours} hour{'s' if hours > 1 else ''} ago"
    elif seconds < 2592000:  # less than a month
        days = int(seconds // 86400)
        return f"{days} day{'s' if days > 1 else ''} ago"
    else:
        # If it's more than a month, return in months
        months = int(seconds // 2592000)
        return f"{months} month{'s' if months > 1 else ''} ago"



@main_bp.route('/')
def index():
    if current_user.is_anonymous:
        return redirect(url_for('main.login'))
    
    days_ago = request.args.get('t', default=3, type=int)
    sort_by = request.args.get('sort_by', default='created_at')
    order = request.args.get('order', default='desc')





    # Validate the `days_ago` parameter
    valid_days = [1, 3, 7, 30]
    if days_ago not in valid_days:
        days_ago = 3

    # Define the date range for recent job postings
    start_time = datetime.utcnow() - timedelta(days=days_ago)

    try:
        # Fetch the current user's job preferences ID
        user_preferences_response = supabase.table('user_job_preferences').select('id').eq('user_id', current_user.id).execute()
        user_preferences = user_preferences_response.data

        if not user_preferences:
            raise ValueError("User preferences not found. Please fill out your job preferences.")

        user_preferences_id = user_preferences[0]['id']

        # Query job postings with associated fit scores
        jobs_response = supabase.table('job_postings').select('*, user_job_fit(fit_score_512)').eq('user_job_fit.user_job_preferences_id', user_preferences_id).gte('created_at', start_time.isoformat()).execute()

        fresh_jobs_list = jobs_response.data  # List of job postings

        # Filter out jobs without a valid fit score
        jobs_to_display = [
            job for job in fresh_jobs_list
            if isinstance(job.get('user_job_fit'), list)
            and len(job['user_job_fit']) > 0
            and job['user_job_fit'][0].get('fit_score_512') is not None
        ]

        # Define the start and end of today as timezone-aware datetime objects
        today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)  # This is just before midnight

        # Filter jobs posted today
        jobs_today = [
            job for job in jobs_to_display
            if datetime.strptime(job['created_at'], '%Y-%m-%dT%H:%M:%S.%f%z') >= today_start
            and datetime.strptime(job['created_at'], '%Y-%m-%dT%H:%M:%S.%f%z') < today_end
        ]
        # Count how many jobs were posted today
        jobs_today_count = len(jobs_today)

        # Step 5: Collect fit scores for percentile calculation
        fit_scores = [job['user_job_fit'][0]['fit_score_512'] for job in jobs_to_display]

        # Step 6: Calculate percentiles for each job's fit score
        for job in jobs_to_display:
            ujf = job.get('user_job_fit', [])[0]
            created_at_time = job.get('created_at', None)
            fit_score = ujf.get('fit_score_512')
            if fit_score is not None:
                job['percentile_512'] = calculate_percentile(fit_scores, fit_score)
            else:
                job['percentile_512'] = None  # This should not occur due to filtering
            
            if created_at_time is not None:
                job['time_ago'] = time_ago(created_at_time)
            else:
                job['time_ago'] = None

                

        # Sort jobs by freshness first (created_at), then by descending fit score
        reverse = (order == 'desc')
        jobs_to_display.sort(key=lambda job: (job.get('created_at'), -job['user_job_fit'][0].get('fit_score_512', 0)), reverse=reverse)

        # Limit jobs to 4
        displayed_jobs = jobs_to_display[:5]

        # Debugging: Log the number of displayed jobs
        logger.debug(f"Displaying {len(displayed_jobs)} jobs sorted by freshness and fit score (days_ago={days_ago}).")

    except Exception as e:
        logger.error(f"Error fetching jobs: {str(e)}")
        flash(f"Error fetching jobs: {str(e)}", 'error')
        return redirect(url_for('main.index'))

    return render_template('index.html',
    user=current_user,
    jobs=displayed_jobs,
    jobs_today=jobs_today_count)

def subscription_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not g.user.is_subscribed:
            flash('This page requires an active subscription.')
            return redirect(url_for('subscription_page'))
        return f(*args, **kwargs)
    return decorated_function

################################
# PAYMENTS AND STRIPE ROUTES
################################

@main_bp.route('/webhook', methods=['POST'])
def webhook():
    event = None
    payload = request.data
    sig_header = request.headers.get('STRIPE_SIGNATURE')

    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, stripe.endpoint_secret
        )
    except ValueError as e:
        # Invalid payload
        raise e
    except stripe.error.SignatureVerificationError as e:
        # Invalid signature
        raise e

    # Handle the event
    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']
        handle_checkout_session(session)
    elif event['type'] == 'customer.subscription.created':
        subscription = event['data']['object']
        handle_subscription_created(subscription)
    elif event['type'] == 'customer.subscription.updated':
        subscription = event['data']['object']
        handle_subscription_updated(subscription)
    elif event['type'] == 'customer.subscription.deleted':
        subscription = event['data']['object']
        handle_subscription_deleted(subscription)
    elif event['type'] in ['invoice.paid', 'invoice.payment_succeeded']:
        invoice = event['data']['object']
        handle_invoice_paid(invoice)
    elif event['type'] == 'invoice.finalized':
        invoice = event['data']['object']
        handle_invoice_finalized(invoice)
    elif event['type'] == 'charge.succeeded':
        charge = event['data']['object']
        handle_charge_succeeded(charge)
    elif event['type'] == 'payment_method.attached':
        payment_method = event['data']['object']
        handle_payment_method_attached(payment_method)
    elif event['type'] == 'customer.created':
        customer = event['data']['object']
        handle_customer_created(customer)
    elif event['type'] == 'customer.updated':
        customer = event['data']['object']
        handle_customer_updated(customer)
    elif event['type'] == 'payment_intent.succeeded':
        payment_intent = event['data']['object']
        handle_payment_intent_succeeded(payment_intent)
    elif event['type'] == 'payment_intent.created':
        payment_intent = event['data']['object']
        handle_payment_intent_created(payment_intent)
    elif event['type'] == 'invoice.created':
        invoice = event['data']['object']
        handle_invoice_created(invoice)
    elif event['type'] == 'invoice.updated':
        invoice = event['data']['object']
        handle_invoice_updated(invoice)
    else:
        print(f"Unhandled event type {event['type']}")

    return jsonify(success=True)

@main_bp.route('/create_checkout_session', methods=['GET'])
@login_required
def create_checkout_session():
    if current_user.is_subscribed:
        flash("You are already subscribed!")
        return redirect(url_for('main.index'))
    stripe_customer_id = get_or_create_stripe_customer(current_user.id)
    try:
        checkout_session = stripe.checkout.Session.create(
            customer=stripe_customer_id,
            payment_method_types=['card'],
            line_items=[
                {
                    'price': stripe.cognibly_subscription_price_id,  # Make sure to set this in your config
                    'quantity': 1,
                },
            ],
            mode='subscription',
            allow_promotion_codes=True,
            success_url=url_for('main.success', _external=True) + '?session_id={CHECKOUT_SESSION_ID}',
            cancel_url=url_for('main.cancel', _external=True),
        )
        return redirect(checkout_session.url, code=303)
    except Exception as e:
        logger.error(f"Error creating checkout session: {str(e)}")
        flash("An error occurred while processing your request. Please try again later.")
        return redirect(url_for('main.index'))

@main_bp.route('/success')
@login_required
def success():
    session_id = request.args.get('session_id')
    if not session_id:
        flash("No session ID provided.", "error")
        return redirect(url_for('main.index'))

    try:
        checkout_session = stripe.checkout.Session.retrieve(session_id)
        
        update_data = {
            'stripe_customer_id': checkout_session.customer,
            'is_subscribed': True,
            'subscription_id': checkout_session.subscription,
            'subscription_status': 'active'
        }
        
        response = supabase.table('profiles').update(update_data).eq('id', current_user.id).execute()
        
        if response.data:
            current_user.stripe_customer_id = checkout_session.customer
            current_user.is_subscribed = True
            current_user.subscription_id = checkout_session.subscription
            current_user.subscription_status = 'active'
            
            flash("Thank you! Subscription successfully created!", "success")
        else:
            flash("There was an issue updating your subscription status. Please contact support.", "error")
    
    except stripe.error.StripeError as e:
        logger.error(f"Stripe error occurred: {str(e)}")
        flash(f"An error occurred with the payment processor: {str(e)}", "error")
    except Exception as e:
        logger.exception("Unexpected error in success route")
        flash("An unexpected error occurred. Our team has been notified.", "error")
    finally:
        # Clear the session ID to prevent repeated processing
        session.pop('checkout_session_id', None)
    
    return redirect(url_for('main.index'))

@main_bp.route('/cancel')
@login_required
def cancel():
    # Log the cancellation (optional)
    logger.info(f"User {current_user.id} cancelled the checkout process.")
    
    # Flash a message to the user
    flash("Your subscription process was cancelled. If you have any questions, please contact support.", "info")
    
    # Redirect to a relevant page (e.g., homepage or a subscription information page)
    return redirect(url_for('main.index'))

@main_bp.route('/cancel-subscription', methods=['GET'])
@login_required
def cancel_subscription():
    logger.debug("Cancel subscription route accessed")
    user_id = current_user.id  # Assuming you have a way to get the current user's ID
    print(f"Cancelling subscription for user ID: {user_id}")
    # Fetch user data from Supabase
    response = supabase.table('profiles').select('subscription_id').eq('id', user_id).execute()
    if len(response.data) == 0 or not response.data[0].get('subscription_id'):
        flash('No active subscription found.', 'warning')
        return redirect(url_for('main.index'))

    subscription_id = response.data[0]['subscription_id']

    try:
        # Cancel the subscription at the end of the current period
        subscription = stripe.Subscription.modify(
            subscription_id,
            cancel_at_period_end=True
        )

        # Update Supabase
        response = supabase.table('profiles').update({
            'subscription_status': 'canceling',
            'cancel_at_period_end': subscription['cancel_at_period_end']
        }).eq('id', user_id).execute()

        flash('Your subscription will be canceled at the end of the current billing period.', 'success')
    except stripe.error.StripeError as e:
        flash(f'An error occurred: {str(e)}', 'error')

    return redirect(url_for('main.index'))

@main_bp.route('/dont-cancel-subscription', methods=['GET'])
@login_required
def dont_cancel_subscription():
    logger.debug("Don't cancel subscription route accessed")
    user_id = current_user.id

    try:
        # Fetch user data from Supabase
        response = supabase.table('profiles').select('subscription_id').eq('id', user_id).execute()
        if len(response.data) == 0 or not response.data[0].get('subscription_id'):
            logger.warning(f"No active subscription found for user {user_id}")
            flash('No active subscription found.', 'warning')
            return redirect(url_for('main.index'))

        subscription_id = response.data[0]['subscription_id']
        logger.info(f"Found subscription {subscription_id} for user {user_id}")

        # Update the Stripe subscription
        subscription = stripe.Subscription.modify(
            subscription_id,
            cancel_at_period_end=False
        )
        logger.info(f"Successfully removed cancellation for Stripe subscription {subscription_id}")

        # Update Supabase
        update_response = supabase.table('profiles').update({
            'subscription_status': 'active',
            'cancel_at_period_end': subscription['cancel_at_period_end']
        }).eq('id', user_id).execute()

        if len(update_response.data) == 0:
            logger.error(f"Supabase update for user {user_id} returned no data")
            flash('An error occurred while updating your subscription status.', 'error')
        else:
            logger.info(f"Successfully updated subscription status to active for user {user_id}")
            flash('Your subscription cancellation has been reversed. Your subscription will continue as normal.', 'success')

    except stripe.error.StripeError as e:
        logger.error(f"Stripe error for user {user_id}: {str(e)}")
        flash(f'An error occurred with Stripe: {str(e)}', 'error')
    except Exception as e:
        logger.error(f"Unexpected error for user {user_id}: {str(e)}")
        flash('An unexpected error occurred. Please try again later.', 'error')

    return redirect(url_for('main.index'))

################################
# STRIPE WEBHOOK HELPER FUNCTIONS
################################

def get_or_create_stripe_customer(user_id):
    # Look up the user's profile
    user_profile = supabase.table('profiles').select('*').eq('id', user_id).execute()
    
    if user_profile.data and user_profile.data[0].get('stripe_customer_id'):
        return user_profile.data[0]['stripe_customer_id']
    
    # If no Stripe customer ID exists, create one
    user = supabase.table('profiles').select('*').eq('id', user_id).execute()
    if not user.data:
        raise ValueError(f"No user found with id {user_id}")
    
    stripe_customer = stripe.Customer.create(email=user.data[0]['email'])
    
    # Update the user's profile with the new Stripe customer ID
    supabase.table('profiles').update({
        'stripe_customer_id': stripe_customer.id
    }).eq('id', user_id).execute()
    
    return stripe_customer.id

def handle_checkout_session(session):
    customer_id = session['customer']
    subscription_id = session['subscription']
    
    # Update user's subscription status in Supabase
    response = supabase.table('profiles').update({
        'is_subscribed': True,
        'subscription_id': subscription_id
    }).eq('stripe_customer_id', customer_id).execute()
    
    if response.data:
        print(f"User subscription updated for customer {customer_id}")
    else:
        print(f"Failed to update user subscription for customer {customer_id}")

def handle_subscription_created(subscription):
    try:
        customer_id = subscription['customer']
        subscription_id = subscription['id']
        cancel_at_period_end = subscription['cancel_at_period_end']
        status = subscription['status']
        
        # Find the user profile with this Stripe customer ID
        user_profile = supabase.table('profiles').select('*').eq('stripe_customer_id', customer_id).execute()
        
        if not user_profile.data:
            print(f"User profile not found for Stripe customer {customer_id}")
            return
        
        user_id = user_profile.data[0]['id']
        
        # Update the user's profile with the subscription information
        response = supabase.table('profiles').update({
            'subscription_id': subscription_id,
            'subscription_status': status,
            'cancel_at_period_end': cancel_at_period_end
        }).eq('id', user_id).execute()
        
        if response.data:
            print(f"Updated subscription for user {user_id}")
        else:
            print(f"Failed to update subscription for user {user_id}")
            print(f"Supabase response: {response}")
    except Exception as e:
        print(f"Error in handle_subscription_created: {str(e)}")
        import traceback
        traceback.print_exc()

def handle_subscription_updated(subscription):
    handle_subscription_created(subscription)

def handle_subscription_deleted(subscription):
    customer_id = subscription['customer']
    
    response = supabase.table('profiles').update({
        'is_subscribed': False,
        'subscription_id': None
    }).eq('stripe_customer_id', customer_id).execute()
    
    if response.data:
        print(f"User subscription deleted for customer {customer_id}")
    else:
        print(f"Failed to delete user subscription for customer {customer_id}")

def handle_invoice_paid(invoice):
    try:
        customer_id = invoice['customer']
        amount_paid = invoice['amount_paid']
        invoice_id = invoice['id']
        
        # Find the user profile with this Stripe customer ID
        user_profile = supabase.table('profiles').select('*').eq('stripe_customer_id', customer_id).execute()
        
        if not user_profile.data:
            print(f"User profile not found for Stripe customer {customer_id}")
            return
        
        user_id = user_profile.data[0]['id']
        
        # Update the user's profile with the payment information
        response = supabase.table('profiles').update({
            'last_payment_amount': amount_paid,
            'last_payment_date': datetime.utcnow().isoformat(),
            'last_invoice_id': invoice_id
        }).eq('id', user_id).execute()
        
        if response.data:
            print(f"Recorded payment for user {user_id}")
        else:
            print(f"Failed to record payment for user {user_id}")
            print(f"Supabase response: {response}")
    except Exception as e:
        print(f"Error in handle_invoice_paid: {str(e)}")
        import traceback
        traceback.print_exc()

def handle_invoice_finalized(invoice):
    customer_id = invoice['customer']
    invoice_id = invoice['id']
    amount_due = invoice['amount_due']
    
    # You might want to log the finalized invoice or update user's billing information
    response = supabase.table('invoices').insert({
        'stripe_customer_id': customer_id,
        'stripe_invoice_id': invoice_id,
        'amount_due': amount_due,
        'status': 'finalized',
        'finalized_at': datetime.utcnow().isoformat()
    }).execute()
    
    if response.data:
        print(f"Invoice finalized for customer {customer_id}")
    else:
        print(f"Failed to record finalized invoice for customer {customer_id}")

def handle_invoice_updated(invoice):
    customer_id = invoice['customer']
    amount_due = invoice['amount_due']
    status = invoice['status']
    print(f"Invoice updated for customer {customer_id}, amount due: {amount_due}, status: {status}")
    # Add any necessary logic here

def handle_charge_succeeded(charge):
    customer_id = charge['customer']
    amount = charge['amount']
    print(f"Charge succeeded for customer {customer_id}, amount: {amount}")
    # Add any necessary logic here

def handle_payment_method_attached(payment_method):
    customer_id = payment_method['customer']
    print(f"Payment method attached for customer {customer_id}")
    # Add any necessary logic here

def handle_customer_created(customer):
    customer_id = customer['id']
    email = customer['email']
    print(f"Customer created: {customer_id}, email: {email}")
    # Add any necessary logic here

def handle_customer_updated(customer):
    customer_id = customer['id']
    print(f"Customer updated: {customer_id}")
    # Add any necessary logic here

def handle_payment_intent_succeeded(payment_intent):
    customer_id = payment_intent['customer']
    amount = payment_intent['amount']
    print(f"Payment intent succeeded for customer {customer_id}, amount: {amount}")
    # Add any necessary logic here

def handle_payment_intent_created(payment_intent):
    customer_id = payment_intent['customer']
    amount = payment_intent['amount']
    print(f"Payment intent created for customer {customer_id}, amount: {amount}")
    # Add any necessary logic here

def handle_invoice_created(invoice):
    customer_id = invoice['customer']
    amount = invoice['amount_due']
    print(f"Invoice created for customer {customer_id}, amount due: {amount}")
    # Add any necessary logic here



@main_bp.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        try:
            # Authenticate user with Supabase Auth
            response = supabase.auth.sign_in_with_password({"email": email, "password": password})
            if response.user:
                user_id = response.user.id

                # Fetch the user's last_login from the profiles table
                user_response = supabase.table('profiles').select('id, last_login', 'is_first_login').eq('id', user_id).execute()

                if user_response.data:
                    user_data = user_response.data[0]
                    last_login = user_data.get('last_login')
                    is_first_login = user_data.get('is_first_login', True)  # Defaults to True if the value is not set

                    if is_first_login:
                        # Set session or pass flag to the front-end to show the modal
                        session['show_profile_modal'] = True


                    # If last_login is more than 7 days ago, trigger Celery task
                    if last_login:
                        # Modify the format to handle timezone offset +%z
                        last_login_time = datetime.strptime(last_login, '%Y-%m-%dT%H:%M:%S.%f%z')

                        # Ensure both datetimes are offset-aware
                        current_time = datetime.utcnow().replace(tzinfo=timezone.utc)  # Make current_time offset-aware

                        # Compare the datetimes
                        if current_time - last_login_time > timedelta(days=7):
                            # Trigger Celery task for this user
                            #celery.send_task('update_all_jobs', args=[user_id])
                            flash('Your account has been inactive for more than 7 days. Please wait while your account gets new jobs to browse.', 'info')

                    # Update last_login to the current timestamp in the profiles table
                    now = datetime.utcnow().replace(tzinfo=timezone.utc)  # Current UTC time with timezone info
                    last_login_str = now.isoformat()  # Convert to ISO 8601 format

                    # Update the last_login field in the profiles table
                    update_response = supabase.table('profiles').update({'last_login': last_login_str}).eq('id', user_id).execute()

                    # Log in the user (store user data in `User` model and log in via Flask-Login)
                    user = User(user_data)
                    login_user(user)
                    flash('Logged in successfully.', 'success')
                    return redirect(url_for('main.index'))
                else:
                    flash('User profile not found. Please contact support.', 'error')
            else:
                flash('Login failed. No user found or incorrect credentials.', 'error')
        except Exception as e:
            flash(f'Login failed: {str(e)}', 'error')

    return render_template('login.html')


@main_bp.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form['email'].strip()
        password = request.form.get('password')
        confirm_password = request.form.get('confirm-password')
        real_name = request.form.get('name')

        # Check if passwords match
        if password != confirm_password:
            flash("Passwords do not match.", "error")
            return redirect(url_for('main.signup'))

        try:
            # Check if email already exists in 'profiles' table
            existing_user = supabase.table('profiles').select('email').eq('email', email).execute()
            if existing_user.data:
                flash('Email already registered. Please log in or use a different email.', 'error')
                return redirect(url_for('main.signup'))

            # Sign up user with Supabase Auth
            response = supabase.auth.sign_up({"email": email, "password": password})

            print(f"Full Supabase Response for Sign Up: {response}")

            user_id = response.user.id  # Get the user ID

            # Insert into profiles table
            profile_response = supabase.table('profiles').insert({
                "id": user_id,
                "email": email
            }).execute()

            job_pref_response = supabase.table('user_job_preferences').insert({
                    "user_id": user_id,
                    "real_name": real_name,
                    "email": email
            }).execute()
            print("Job preferences inserted:", job_pref_response)


            flash('Signup successful. Please check your email to verify your account.', 'success')
            return redirect(url_for('main.login'))

        except Exception as e:
            import traceback
            traceback.print_exc()  # Logs full error in the console for debugging
            flash(f'Signup failed: {str(e)}', 'error')

    return render_template('signup.html')


# Restored Supabase logout route
@main_bp.route('/logout')
@login_required
def logout():
    try:
        # Sign out from Supabase
        supabase.auth.sign_out()
        
        # Log out the user from the Flask application
        logout_user()
        
        flash('You have been successfully logged out.', 'success')
    except Exception as e:
        flash(f'Logout failed: {str(e)}', 'error')
    
    return redirect(url_for('main.index'))

def calculate_percentile(data, value):
    """Calculates the non-reversed percentile of a value within a dataset."""
    if not data:
        return float('nan')  # Handle empty data gracefully

    sorted_data = sorted(data)
    n = len(sorted_data)
    rank = sum(1 for x in sorted_data if x < value)
    percentile = round((rank / n) * 100)
    #print(f"Calculated percentile for {value}: {str(percentile)}")
    return percentile

# @main_bp.route('/pricing')
# def pricing():
#     # Get the user's current subscription status if logged in
#     is_subscribed = current_user.is_subscribed if not current_user.is_anonymous else False
    
#     # Define the pricing tiers
#     pricing_tiers = {
#         'free': {
#             'name': 'Free',
#             'price': '0',
#             'period': 'forever',
#             'features': [
#                 '10 daily job recommendations',
#                 'Indeed job board only',
#                 'Basic job matching',
#                 '7-day job history',
#                 'Basic search filters',
#                 'Single location search',
#                 'Up to 3 industry preferences',
#                 'Up to 2 role types',
#                 'Basic salary range matching'
#             ]
#         },
#         'premium': {
#             'name': 'Premium',
#             'price': '49.99',
#             'period': 'month',
#             'features': [
#                 'Unlimited job recommendations',
#                 'Multiple job boards',
#                 'Advanced AI matching algorithm',
#                 'Real-time job alerts',
#                 'Unlimited job history',
#                 'AI cover letter generation',
#                 'Auto-tailored resumes',
#                 'Detailed job fit analysis',
#                 'Skills gap analysis',
#                 'AI interview preparation',
#                 'Multiple location search',
#                 'Unlimited industry preferences',
#                 'Advanced salary insights',
#                 'Priority support',
#                 'Weekly job market reports'
#             ]
#         }
#     }
    
#     return render_template('pricing.html', 
#                          pricing_tiers=pricing_tiers, 
#                          is_subscribed=is_subscribed)

@main_bp.route('/pricing')
def pricing():
    # Get the user's current subscription status if logged in
    is_subscribed = current_user.is_subscribed if not current_user.is_anonymous else False
    return render_template('pricing.html', is_subscribed=is_subscribed)

@main_bp.route('/manage_subscription')
def manage_subscription():
    if current_user.is_anonymous:
        is_subscribed = False
        next_payment_amount = None
        next_payment_date = None
    else:
        is_subscribed = current_user.is_subscribed
        try:
            response = supabase.table('profiles').select('next_payment_amount', 'next_payment_date') \
                .eq('id', current_user.id).execute()
            user_data = response.data[0] if response.data else {}
            next_payment_amount = user_data.get('next_payment_amount')
            next_payment_date = user_data.get('next_payment_date')

            # Format next payment date
            if next_payment_date:
                next_payment_date = parser.parse(next_payment_date).strftime('%Y-%m-%d')

            # Format next payment amount as $##.##
            if next_payment_amount is not None:
                # Ensure next_payment_amount is a number, then format as $##.##
                next_payment_amount = f"${float(next_payment_amount):,.2f}"

        except Exception as e:
            logger.exception(f"Error fetching subscription details for user {current_user.id}: {str(e)}")
            next_payment_amount = None
            next_payment_date = None

    return render_template(
        'subscription.html',
        is_subscribed=is_subscribed,
        next_payment_amount=next_payment_amount,
        next_payment_date=next_payment_date
    )

@main_bp.route('/calculate_job_diff')
def calculate_job_diff():
    key = request.args.get('key')
    print(key)
    if key == "oiajsfo123jcfneiaiej23oj2oj3faasd":
        job_id = request.args.get('job_id')
        profile_id = request.args.get('profile_id')
        fit = calculate_user_job_fit(profile_id,job_id)
        return ("Success.")
    else: return ("Error.")

def get_start_time(freshness):
    """Returns a UTC datetime based on freshness filter."""
    if freshness == 'day':
        return datetime.now(timezone.utc) - timedelta(days=1)
    elif freshness == 'week':
        return datetime.now(timezone.utc) - timedelta(weeks=1)
    elif freshness == 'month':
        return datetime.now(timezone.utc) - timedelta(weeks=4)  # Approximate 4 weeks
    else:  # 'all'
        return datetime(2000, 1, 1, tzinfo=timezone.utc)  # Ensure it's in UTC


@main_bp.route('/jobs/items', methods=['GET'])
@login_required
def job_items():
    try:
        # Parse JSON body for offset, limit, freshness, sort column, and sort order
        request_data = request.args
        sort_by = request_data.get('sort_by', 'fit_score_512')  # Default sorting by fit_score_512
        sort_order = request_data.get('sort_order', 'desc')  # Default sorting order: descending
        freshness = request_data.get('freshness', 'day')  # Freshness can be 'day', 'week', 'month', or 'all'. Set default to 'day'.

        if current_user.is_subscribed:
            limit = 1000000
        else:
            limit = 10  # Cap limit at 10 if not subscribed

        if sort_order not in ['asc', 'desc']:
            return jsonify({"status": "error", "message": "Invalid sort order. Must be 'asc' or 'desc'."}), 400
        
        if freshness not in ['day', 'week', 'month', 'all']:
            return jsonify({"status": "error", "message": "Invalid freshness parameter. Must be 'day', 'week', 'month', or 'all'."}), 400

        # Fetch the current user's job preferences ID
        user_preferences_response = supabase.table('user_job_preferences').select('id').eq('user_id', current_user.id).execute()
        user_preferences = user_preferences_response.data

        if not user_preferences:
            return jsonify({"status": "error", "message": "User preferences not found. Please fill out your job preferences."}), 404

        user_preferences_id = user_preferences[0]['id']


        # Get the correct UTC start time
        start_time = get_start_time(freshness)

        # Convert to the exact format '%Y-%m-%dT%H:%M:%S.%f%z'
        formatted_start_time = start_time.strftime('%Y-%m-%dT%H:%M:%S.%f%z')

        # Ensure UTC is correctly formatted (Supabase expects `Z` instead of `+0000`)
        formatted_start_time = formatted_start_time[:-2] + ':' + formatted_start_time[-2:]
        
        # Step 3: Query job postings with associated fit scores
        jobs_response = supabase.table('job_postings').select('*, user_job_fit(fit_score_512)') \
            .eq('user_job_fit.user_job_preferences_id', user_preferences_id) \
            .gte('created_at', str(formatted_start_time))

        # Execute the query and fetch the response data
        jobs_response_data = jobs_response.execute().data  # Call execute() to retrieve the data
        print(f"Jobs Response Data: {jobs_response_data}")
        # Step 4: Filter out jobs without a valid fit score
        jobs_to_display = [
            job for job in jobs_response_data
            if isinstance(job.get('user_job_fit'), list)
            and len(job['user_job_fit']) > 0
            and job['user_job_fit'][0].get('fit_score_512') is not None
        ]

        # Debugging: Log the number of jobs after filtering
        logger.debug(f"Total jobs after filtering: {len(jobs_to_display)}")

        # Step 5: Sorting the jobs based on the fit_score_512
        jobs_to_display.sort(
            key=lambda x: x['user_job_fit'][0]['fit_score_512'], 
            reverse=(sort_order == 'desc')  # Sort descending if order is 'desc', else ascending
        )


        # Step 7: Collect fit scores for percentile calculation
        fit_scores = [job['user_job_fit'][0]['fit_score_512'] for job in jobs_to_display]
    
        # Step 8: Calculate percentiles for each job's fit score
        for job in jobs_to_display:
            ujf = job.get('user_job_fit', [])[0]
            fit_score = ujf.get('fit_score_512')
            if fit_score is not None:
                job['percentile_512'] = calculate_percentile(fit_scores, fit_score)
            else:
                job['percentile_512'] = None  # This should not occur due to filtering

        # Step 9: Return jobs with pagination
        total_count = len(jobs_to_display)
        response = {
            "data": jobs_to_display,
            "pagination": {
                "limit": limit,
                "total_count": total_count
            },
            "status": "success",
        }

        # Return JSON response
        return jsonify(response)

    except Exception as e:
        logger.exception(f"Error occurred in /api/items: {str(e)}")
        return jsonify({"status": "error", "message": "An error occurred while fetching jobs."}), 500

@main_bp.route('/jobs')
@login_required
def jobs():
    # Extract query parameters with default values
    is_subscribed = current_user.is_subscribed if not current_user.is_anonymous else False
    per_page = request.args.get('perPage', default=10, type=int)
    page = request.args.get('page', default=1, type=int)
    freshness = request.args.get('freshness', default='day')
    if not is_subscribed:
        per_page = 10  # Non-subscribed users can only view up to 10 jobs
    if page > 1:
        page = 1  # Non-subscribed users can only access the first page
    if page < 1:
        page = 1
    if per_page < 1:
        per_page = 10
    days_ago = request.args.get('t', default=3, type=int)
    sort_by = request.args.get('sort_by', default='percentile')
    order = request.args.get('order', default='desc')

    # Calculate pagination indices
    start = (page - 1) * per_page
    end = start + per_page  # Slicing is exclusive of the end index

    try:
        # Step 1: Fetch the current user's job preferences ID
        user_preferences_response = supabase.table('user_job_preferences').select('id').eq('user_id', current_user.id).execute()
        user_preferences = user_preferences_response.data

        if not user_preferences:
            raise ValueError("User preferences not found. Please fill out your job preferences.")

        user_preferences_id = user_preferences[0]['id']

        # Step 2: Define the date range for recent job postings
        start_time = datetime.utcnow() - timedelta(days=days_ago)

        # Step 3: Query job postings with associated fit scores
        jobs_response = supabase.table('job_postings').select('*, user_job_fit(fit_score_512)').eq('user_job_fit.user_job_preferences_id', user_preferences_id).gte('created_at', start_time.isoformat()).execute()
        
        fresh_jobs_list = jobs_response.data  # List of job postings

        # Step 4: Filter out jobs without a valid fit score
        # This ensures only jobs with a fit_score_512 are processed and displayed
        jobs_to_display = [
            job for job in fresh_jobs_list
            if isinstance(job.get('user_job_fit'), list)
            and len(job['user_job_fit']) > 0
            and job['user_job_fit'][0].get('fit_score_512') is not None
        ]

        # Define the start and end of today as timezone-aware datetime objects
        today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)  # This is just before midnight

        # Filter jobs posted today
        jobs_today = [
            job for job in jobs_to_display
            if datetime.strptime(job['created_at'], '%Y-%m-%dT%H:%M:%S.%f%z') >= today_start
            and datetime.strptime(job['created_at'], '%Y-%m-%dT%H:%M:%S.%f%z') < today_end
        ]
        # Count how many jobs were posted today
        jobs_today_count = len(jobs_today)


        # Debugging: Log the number of jobs after filtering
        logger.debug(f"Total jobs after filtering: {len(jobs_to_display)}")

        # Step 5: Collect fit scores for percentile calculation
        fit_scores = [job['user_job_fit'][0]['fit_score_512'] for job in jobs_to_display]

        # Step 6: Calculate percentiles for each job's fit score
        for job in jobs_to_display:
            ujf = job.get('user_job_fit', [])[0]
            fit_score = ujf.get('fit_score_512')
            if fit_score is not None:
                job['percentile_512'] = calculate_percentile(fit_scores, fit_score)
            else:
                job['percentile_512'] = None  # This should not occur due to filtering

        # Step 7: Define the sorting key based on the 'sort_by' parameter
        def get_sort_key(job):
            if sort_by == 'percentile':
                percentile = job.get('percentile_512')
                # Jobs without a percentile are considered lower in the sort order
                is_none = percentile is None
                value = percentile if percentile is not None else 0
                return (is_none, value)
            elif sort_by in ['company', 'job_title', 'location']:
                field_value = job.get(sort_by, '')
                is_empty = not field_value
                value = field_value.lower()
                return (is_empty, value)
            else:
                # Default sorting by percentile if an unknown sort_by parameter is provided
                percentile = job.get('percentile_512')
                is_none = percentile is None
                value = percentile if percentile is not None else 0
                return (is_none, value)

        # Determine if the sorting should be in descending order
        reverse = (order == 'desc')

        # Step 8: Sort the jobs based on the defined key function
        jobs_to_display.sort(key=get_sort_key, reverse=reverse)

        # Step 9: Implement pagination after sorting
        job_count = len(jobs_to_display)
        total_pages = ceil(job_count / per_page) if per_page else 1
        if not is_subscribed:
            displayed_jobs = jobs_to_display[start:end][:10]
            total_pages = 1
        else:
            displayed_jobs = jobs_to_display[start:end]

        # Debugging: Log pagination details
        logger.debug(f"Displaying page {page} of {total_pages} with {per_page} jobs per page.")


        has_jobs = len(jobs_to_display) > 0

        # Step 10: Render the 'jobs.html' template with the paginated job listings
        return render_template(
            'jobs.html',
            jobs=displayed_jobs,
            per_page=per_page,
            page=page,
            total_pages=total_pages,
            sort_by=sort_by,
            order=order,
            days_ago=days_ago,
            is_subscribed=is_subscribed,
            jobs_today=jobs_today_count,
            has_jobs=has_jobs,
            freshness=freshness
        )

    except ValueError as e:
        # Handle known errors (e.g., missing user preferences)
        flash(str(e), 'error')
        logger.error(f"ValueError in /jobs route: {str(e)}")
        return redirect(url_for('main.index'))  # Redirect to home or appropriate page
    except Exception as e:
        # Handle unexpected errors
        flash("An unexpected error occurred while fetching jobs.", 'error')
        logger.exception(f"Unexpected error in /jobs route: {str(e)}")
        return redirect(url_for('main.index'))  # Redirect to home or appropriate page

# Fetch user preferences
def get_user_preferences(user_id):
    response = supabase.table('user_job_preferences').select('*').eq('user_id', user_id).execute()
    return response.data[0] if response.data else {}

# Update user preferences
def update_user_preferences(user_id, category, tags):
    response = supabase.table('user_job_preferences').update({category: tags}).eq('user_id', user_id).execute()
    return response

# Suggestions data
suggestions = {
    "preferred_industries":[ 
    # Technology and IT
    "Software Development", "Cloud Computing", "Artificial Intelligence", "Machine Learning",
    "Cybersecurity", "Data Analytics", "Internet of Things (IoT)", "Blockchain",
    "Quantum Computing", "Robotics", "Virtual Reality", "Augmented Reality",
    "5G Technology", "Edge Computing", "DevOps", "IT Consulting",

    # Healthcare and Life Sciences
    "Pharmaceuticals", "Biotechnology", "Medical Devices", "Healthcare IT",
    "Telemedicine", "Genomics", "Personalized Medicine", "Mental Health Services",
    "Elder Care", "Veterinary Medicine", "Dental Care", "Physical Therapy",
    "Nutrition and Wellness", "Health Insurance", "Medical Research", "Public Health",

    # Finance and Banking
    "Commercial Banking", "Investment Banking", "Asset Management", "Hedge Funds",
    "Venture Capital", "Private Equity", "Insurance", "Financial Technology (FinTech)",
    "Cryptocurrency", "Personal Finance", "Accounting", "Tax Services",
    "Real Estate Investment", "Mortgage Lending", "Credit Services", "Financial Consulting",

    # Education and Training
    "K-12 Education", "Higher Education", "Online Education", "EdTech",
    "Corporate Training", "Language Learning", "Special Education", "Early Childhood Education",
    "Vocational Training", "Test Preparation", "Educational Consulting", "Tutoring Services",
    "Adult Education", "STEM Education", "Art Education", "Music Education",

    # Manufacturing and Industry
    "Automotive Manufacturing", "Aerospace Manufacturing", "Electronics Manufacturing",
    "Textile Manufacturing", "Food Processing", "Chemical Manufacturing",
    "Pharmaceutical Manufacturing", "Industrial Automation", "3D Printing",
    "Packaging", "Metalworking", "Plastics Manufacturing", "Paper and Pulp",
    "Machinery Manufacturing", "Furniture Manufacturing", "Toy Manufacturing",

    # Retail and E-commerce
    "Online Retail", "Brick-and-Mortar Retail", "Grocery", "Fashion and Apparel",
    "Luxury Goods", "Consumer Electronics", "Home Improvement", "Sporting Goods",
    "Beauty and Cosmetics", "Pet Supplies", "Jewelry", "Books and Media",
    "Furniture and Home Decor", "Office Supplies", "Automotive Retail", "Specialty Foods",

    # Hospitality and Tourism
    "Hotels and Resorts", "Restaurants", "Fast Food", "Catering",
    "Travel Agencies", "Airlines", "Cruise Lines", "Theme Parks",
    "Casinos and Gaming", "Event Planning", "Tour Operators", "Vacation Rentals",
    "Spa and Wellness Centers", "Timeshare", "Eco-Tourism", "Cultural Tourism",

    # Energy and Utilities
    "Oil and Gas", "Renewable Energy", "Solar Power", "Wind Power",
    "Hydroelectric Power", "Nuclear Energy", "Energy Storage", "Smart Grid Technology",
    "Waste Management", "Water Treatment", "Natural Gas Distribution", "Electric Utilities",
    "Energy Efficiency", "Geothermal Energy", "Biomass Energy", "Hydrogen Fuel Cells",

    # Telecommunications
    "Wireless Carriers", "Broadband Providers", "Satellite Communications", "Fiber Optics",
    "Telecom Equipment", "VoIP Services", "Network Infrastructure", "Telecom Software",
    "Mobile Virtual Network Operators", "Unified Communications", "Telecom Consulting", "Data Centers",

    # Automotive
    "Car Manufacturing", "Electric Vehicles", "Autonomous Vehicles", "Auto Parts Manufacturing",
    "Car Dealerships", "Auto Repair and Maintenance", "Fleet Management", "Car Rental",
    "Automotive Design", "Motorcycle Manufacturing", "Truck Manufacturing", "Automotive Software",

    # Aerospace and Defense
    "Aircraft Manufacturing", "Space Technology", "Satellite Systems", "Defense Contracting",
    "Missile Systems", "Military Vehicles", "Avionics", "Drone Technology",
    "Air Traffic Control Systems", "Aircraft Maintenance", "Space Exploration", "Rocket Propulsion",

    # Agriculture and Farming
    "Crop Farming", "Livestock Farming", "Organic Farming", "Precision Agriculture",
    "Aquaculture", "Forestry", "Agricultural Biotechnology", "Farm Equipment",
    "Seed Technology", "Pesticides and Fertilizers", "Vertical Farming", "Hydroponics",
    "Agricultural Drones", "Food Safety", "Sustainable Agriculture", "Agritourism",

    # Construction and Real Estate
    "Residential Construction", "Commercial Construction", "Infrastructure Development",
    "Architecture", "Civil Engineering", "Real Estate Development", "Property Management",
    "Interior Design", "Landscape Architecture", "Building Materials", "Smart Home Technology",
    "Green Building", "Facilities Management", "Urban Planning", "Real Estate Investment Trusts",

    # Entertainment and Media
    "Film Production", "Television Broadcasting", "Streaming Services", "Music Industry",
    "Video Game Development", "Publishing", "Advertising", "Public Relations",
    "Social Media", "News Media", "Animation", "Podcasting",
    "Live Events", "Sports Entertainment", "Radio Broadcasting", "Digital Marketing",

    # Transportation and Logistics
    "Trucking", "Rail Transport", "Air Cargo", "Maritime Shipping",
    "Logistics Software", "Warehousing", "Supply Chain Management", "Last-Mile Delivery",
    "Freight Forwarding", "Autonomous Logistics", "Cold Chain Logistics", "Reverse Logistics",
    "Intermodal Transportation", "Logistics Consulting", "Postal Services", "Courier Services",

    # Environmental Services
    "Environmental Consulting", "Pollution Control", "Recycling", "Green Technology",
    "Climate Change Mitigation", "Conservation", "Sustainable Development", "Ecological Restoration",
    "Environmental Impact Assessment", "Hazardous Waste Management", "Air Quality Management", "Water Conservation",

    # Professional Services
    "Legal Services", "Management Consulting", "Human Resources", "Recruitment",
    "Marketing Services", "Graphic Design", "Engineering Services", "Market Research",
    "Business Process Outsourcing", "Translation Services", "Data Entry Services", "Transcription Services",

    # Non-Profit and Social Services
    "Charitable Organizations", "Social Advocacy", "Community Development", "Disaster Relief",
    "International Aid", "Human Rights", "Animal Welfare", "Environmental Conservation",
    "Arts and Culture", "Education Foundations", "Healthcare Foundations", "Religious Organizations"
],

    "preferred_roles_responsibilities": [
    # Management Roles
    "Chief Executive Officer (CEO)", "Chief Financial Officer (CFO)", "Chief Operating Officer (COO)",
    "Chief Technology Officer (CTO)", "Chief Marketing Officer (CMO)", "Chief Human Resources Officer (CHRO)",
    "Director of Operations", "Project Manager", "Program Manager", "Department Manager",
    "Team Leader", "Supervisor",

    # Finance and Accounting
    "Financial Analyst", "Accountant", "Auditor", "Budget Analyst", "Tax Specialist",
    "Investment Banker", "Financial Planner", "Risk Manager", "Actuary",

    # Technology and IT
    "Software Engineer", "Full Stack Developer", "Front-end Developer", "Back-end Developer",
    "Data Scientist", "Database Administrator", "Systems Administrator", "Network Engineer",
    "Cloud Architect", "DevOps Engineer", "Information Security Analyst", "UI/UX Designer",
    "Machine Learning Engineer", "Artificial Intelligence Specialist", "QA Engineer",

    # Sales and Marketing
    "Sales Representative", "Account Manager", "Business Development Manager",
    "Marketing Specialist", "Digital Marketing Manager", "Content Marketing Manager",
    "Brand Manager", "Product Marketing Manager", "SEO Specialist", "Social Media Manager",
    "Public Relations Specialist", "Market Research Analyst",

    # Human Resources
    "HR Manager", "Recruiter", "Training and Development Specialist",
    "Compensation and Benefits Analyst", "Employee Relations Specialist",
    "Talent Acquisition Manager", "HR Business Partner", "Diversity and Inclusion Specialist",

    # Operations and Logistics
    "Operations Manager", "Supply Chain Manager", "Logistics Coordinator",
    "Procurement Specialist", "Inventory Manager", "Quality Assurance Manager",
    "Facilities Manager", "Production Planner", "Process Improvement Specialist",

    # Customer Service
    "Customer Service Representative", "Customer Success Manager",
    "Technical Support Specialist", "Client Relations Manager",

    # Research and Development
    "Research Scientist", "Product Developer", "R&D Manager",
    "Innovation Specialist", "Patent Specialist",

    # Legal
    "Corporate Lawyer", "Legal Counsel", "Compliance Officer",
    "Paralegal", "Intellectual Property Specialist",

    # Healthcare
    "Physician", "Nurse", "Pharmacist", "Medical Researcher",
    "Healthcare Administrator", "Physical Therapist", "Occupational Therapist",

    # Education
    "Teacher", "Professor", "Education Administrator", "Curriculum Developer",
    "Instructional Designer", "School Counselor", "Special Education Specialist",

    # Creative and Design
    "Graphic Designer", "Art Director", "Copywriter", "Video Editor",
    "UX/UI Designer", "Product Designer", "Industrial Designer",

    # Engineering
    "Mechanical Engineer", "Civil Engineer", "Electrical Engineer",
    "Chemical Engineer", "Aerospace Engineer", "Environmental Engineer",

    # Consulting
    "Management Consultant", "Strategy Consultant", "IT Consultant",
    "Financial Consultant", "Human Resources Consultant",

    # Data and Analytics
    "Data Analyst", "Business Intelligence Analyst", "Data Engineer",
    "Statistician", "Operations Research Analyst",

    # Project Management
    "Project Coordinator", "Scrum Master", "Agile Coach",
    "Program Director", "Portfolio Manager",

    # Communications
    "Communications Specialist", "Technical Writer", "Translator",
    "Interpreter", "Copyeditor",

    # Specialized Roles
    "Sustainability Officer", "Diversity and Inclusion Manager",
    "Change Management Specialist", "Crisis Management Specialist",
    "Innovation Manager", "Knowledge Management Specialist"
],

    "preferred_locations": [    "New York, NY", "Los Angeles, CA", "Chicago, IL", "Houston, TX", "Phoenix, AZ",
    "Philadelphia, PA", "San Antonio, TX", "San Diego, CA", "Dallas, TX", "San Jose, CA",
    "Austin, TX", "Jacksonville, FL", "San Francisco, CA", "Columbus, OH", "Fort Worth, TX",
    "Indianapolis, IN", "Charlotte, NC", "Seattle, WA", "Denver, CO", "El Paso, TX",
    "Washington, DC", "Boston, MA", "Detroit, MI", "Nashville, TN", "Oklahoma City, OK",
    "Portland, OR", "Las Vegas, NV", "Louisville, KY", "Baltimore, MD", "Milwaukee, WI",
    "Albuquerque, NM", "Tucson, AZ", "Fresno, CA", "Sacramento, CA", "Long Beach, CA",
    "Kansas City, MO", "Mesa, AZ", "Virginia Beach, VA", "Atlanta, GA", "Colorado Springs, CO",
    "Omaha, NE", "Raleigh, NC", "Miami, FL", "Cleveland, OH", "Tulsa, OK",
    "Oakland, CA", "Minneapolis, MN", "Wichita, KS", "New Orleans, LA", "Arlington, TX",
    "Bakersfield, CA", "Tampa, FL", "Honolulu, HI", "Anaheim, CA", "Santa Ana, CA",
    "Corpus Christi, TX", "Riverside, CA", "St. Louis, MO", "Pittsburgh, PA", "Greensboro, NC",
    "Lincoln, NE", "Anchorage, AK", "Plano, TX", "Orlando, FL", "Irvine, CA",
    "Laredo, TX", "Chula Vista, CA", "Durham, NC", "Jersey City, NJ", "Fort Wayne, IN",
    "St. Petersburg, FL", "Chandler, AZ", "Lubbock, TX", "Madison, WI", "Gilbert, AZ",
    "Reno, NV", "Hialeah, FL", "Baton Rouge, LA", "Richmond, VA", "Boise, ID",
    "San Bernardino, CA", "Spokane, WA", "Des Moines, IA", "Modesto, CA", "Fremont, CA",
    "Santa Clarita, CA", "Mobile, AL", "Oxnard, CA", "Moreno Valley, CA", "Huntington Beach, CA",
    "Aurora, CO", "Columbia, SC", "Grand Rapids, MI", "Salt Lake City, UT", "Tallahassee, FL",
    "Overland Park, KS", "Knoxville, TN", "Worcester, MA", "Newport News, VA", "Brownsville, TX",
    "Santa Rosa, CA", "Vancouver, WA", "Fort Lauderdale, FL", "Sioux Falls, SD", "Ontario, CA"]
}

# Route to get suggestions
@main_bp.route('/get-suggestions/<category>', methods=['GET'])
def get_suggestions(category):
    query = request.args.get('query', '').lower()
    if query:
        filtered = [s for s in suggestions.get(category, []) if query in s.lower()]
        return jsonify(filtered)
    return jsonify([])

@main_bp.route('/get-tags/<category>', methods=['GET'])
def get_tags(category):
    user_id = current_user.id  # Fetch the logged-in user's ID
    if not user_id:
        return jsonify({"error": "User not authenticated"}), 401

    # Fetch preferences from the database
    preferences = get_user_preferences(user_id)
    tags = preferences.get(category, '').split(',') if preferences.get(category) else []

    return jsonify(tags)


# Route to add a tag
@main_bp.route('/add-tag/<category>', methods=['POST'])
def add_tag(category):
    user_id = current_user.id  # Retrieve the logged-in user's ID
    tag = request.json.get('tag')

    # Fetch current preferences
    preferences = get_user_preferences(user_id)
    tags = preferences.get(category, '').split(',') if preferences.get(category) else []
    
    if tag and tag not in tags:
        tags.append(tag)
        # Update database
        update_user_preferences(user_id, category, ','.join(tags))
    
    return jsonify({"message": f"Tag added to {category}!", "data": tags})

# Route to remove a tag
@main_bp.route('/remove-tag/<category>', methods=['DELETE'])
def remove_tag(category):
    user_id = current_user.id  # Retrieve the logged-in user's ID
    tag = request.json.get('tag')

    # Fetch current preferences
    preferences = get_user_preferences(user_id)
    tags = preferences.get(category, '').split(',') if preferences.get(category) else []
    
    if tag in tags:
        tags.remove(tag)
        # Update database
        update_user_preferences(user_id, category, ','.join(tags))
    
    return jsonify({"message": f"Tag removed from {category}!", "data": tags})

# Route to clear tags for a category
@main_bp.route('/clear-tags/<category>', methods=['POST'])
def clear_tags(category):
    user_id = request.json.get('user_id')  # Assume the front-end sends the user ID

    # Update database
    update_user_preferences(user_id, category, '')
    
    return jsonify({"message": f"All tags cleared for {category}!", "data": []})

@main_bp.route('/update-salary-range', methods=['POST'])
def update_salary_range():
    # Get the salary range from the request
    data = request.get_json()
    salary_range = data.get('salary_range')

    # Ensure the salary range is not None or empty
    if not salary_range:
        return jsonify({"error": "Salary range is required"}), 400

    # Update the user's job preferences table with the new salary range
    try:
        response = supabase.table('user_job_preferences').update({
            'expected_salary_range': salary_range
        }).eq('user_id', current_user.id).execute()

        # Check if the update was successful
        if response.data:
            return jsonify({"message": "Salary range updated successfully"}), 200
        else:
            return jsonify({"error": "Failed to update salary range"}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Routes for preferred industries
@main_bp.route('/preferred/industries', methods=['GET', 'POST', 'DELETE'])
def manage_preferred_industries():
    try:
        if request.method == 'GET':
            # Fetch current preferred industries
            response = supabase.table('user_job_preferences').select('preferred_industries').eq('user_id', current_user.id).single().execute()
            preferred_industries_str = response.data.get('preferred_industries', '')

            # Convert comma-separated string to list
            preferred_industries = preferred_industries_str.split(',') if preferred_industries_str else []
            return jsonify(preferred_industries), 200

        elif request.method == 'POST':
            data = request.get_json()
            new_industry = data.get('industry')

            if not new_industry:
                return jsonify({"error": "Industry is required"}), 400

            # Fetch current preferred industries
            response = supabase.table('user_job_preferences').select('preferred_industries').eq('user_id', current_user.id).single().execute()
            preferred_industries_str = response.data.get('preferred_industries', '')

            # Convert to list for manipulation
            preferred_industries = preferred_industries_str.split(',') if preferred_industries_str else []

            if new_industry in preferred_industries:
                return jsonify({"error": "Industry already added"}), 400

            # Add new industry
            preferred_industries.append(new_industry)

            # Convert back to comma-separated string
            updated_industries_str = ','.join(preferred_industries)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_industries': updated_industries_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Industry added successfully", "industries": preferred_industries}), 200

        elif request.method == 'DELETE':
            data = request.get_json()
            industry_to_remove = data.get('industry')

            if not industry_to_remove:
                return jsonify({"error": "Industry is required for deletion"}), 400

            # Fetch current preferred industries
            response = supabase.table('user_job_preferences').select('preferred_industries').eq('user_id', current_user.id).single().execute()
            preferred_industries_str = response.data.get('preferred_industries', '')

            # Convert to list for manipulation
            preferred_industries = preferred_industries_str.split(',') if preferred_industries_str else []

            if industry_to_remove not in preferred_industries:
                return jsonify({"error": "Industry not found"}), 400

            # Remove industry
            preferred_industries.remove(industry_to_remove)

            # Convert back to comma-separated string
            updated_industries_str = ','.join(preferred_industries)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_industries': updated_industries_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Industry removed successfully", "industries": preferred_industries}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Routes for preferred roles and responsibilities
@main_bp.route('/preferred/roles', methods=['GET', 'POST', 'DELETE'])
def manage_preferred_roles_responsibilities():
    try:
        if request.method == 'GET':
            # Fetch current preferred roles and responsibilities
            response = supabase.table('user_job_preferences').select('preferred_roles_responsibilities').eq('user_id', current_user.id).single().execute()
            preferred_roles_responsibilities_str = response.data.get('preferred_roles_responsibilities', '')

            # Convert comma-separated string to list
            preferred_roles_responsibilities = preferred_roles_responsibilities_str.split(',') if preferred_roles_responsibilities_str else []
            return jsonify(preferred_roles_responsibilities), 200

        elif request.method == 'POST':
            data = request.get_json()
            new_role_responsibility = data.get('role_responsibility')

            if not new_role_responsibility:
                return jsonify({"error": "Role or Responsibility is required"}), 400

            # Fetch current preferred roles and responsibilities
            response = supabase.table('user_job_preferences').select('preferred_roles_responsibilities').eq('user_id', current_user.id).single().execute()
            preferred_roles_responsibilities_str = response.data.get('preferred_roles_responsibilities', '')

            # Convert to list for manipulation
            preferred_roles_responsibilities = preferred_roles_responsibilities_str.split(',') if preferred_roles_responsibilities_str else []

            if new_role_responsibility in preferred_roles_responsibilities:
                return jsonify({"error": "Role or Responsibility already added"}), 400

            # Add new role/responsibility
            preferred_roles_responsibilities.append(new_role_responsibility)

            # Convert back to comma-separated string
            updated_roles_responsibilities_str = ','.join(preferred_roles_responsibilities)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_roles_responsibilities': updated_roles_responsibilities_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Role or Responsibility added successfully", "roles_responsibilities": preferred_roles_responsibilities}), 200

        elif request.method == 'DELETE':
            data = request.get_json()
            role_responsibility_to_remove = data.get('role_responsibility')

            if not role_responsibility_to_remove:
                return jsonify({"error": "Role or Responsibility is required for deletion"}), 400

            # Fetch current preferred roles and responsibilities
            response = supabase.table('user_job_preferences').select('preferred_roles_responsibilities').eq('user_id', current_user.id).single().execute()
            preferred_roles_responsibilities_str = response.data.get('preferred_roles_responsibilities', '')

            # Convert to list for manipulation
            preferred_roles_responsibilities = preferred_roles_responsibilities_str.split(',') if preferred_roles_responsibilities_str else []

            if role_responsibility_to_remove not in preferred_roles_responsibilities:
                return jsonify({"error": "Role or Responsibility not found"}), 400

            # Remove role/responsibility
            preferred_roles_responsibilities.remove(role_responsibility_to_remove)

            # Convert back to comma-separated string
            updated_roles_responsibilities_str = ','.join(preferred_roles_responsibilities)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_roles_responsibilities': updated_roles_responsibilities_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Role or Responsibility removed successfully", "roles_responsibilities": preferred_roles_responsibilities}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@main_bp.route('/preferred/locations', methods=['GET', 'POST', 'DELETE'])
def manage_preferred_locations():
    try:
        if request.method == 'GET':
            # Fetch current preferred locations
            response = supabase.table('user_job_preferences').select('preferred_locations').eq('user_id', current_user.id).single().execute()
            preferred_locations_str = response.data.get('preferred_locations', '')

            # Convert comma-separated string to list
            preferred_locations = preferred_locations_str.split(',') if preferred_locations_str else []
            return jsonify(preferred_locations), 200

        elif request.method == 'POST':
            data = request.get_json()
            new_location = data.get('location')

            if not new_location:
                return jsonify({"error": "Location is required"}), 400

            # Fetch current preferred locations
            response = supabase.table('user_job_preferences').select('preferred_locations').eq('user_id', current_user.id).single().execute()
            preferred_locations_str = response.data.get('preferred_locations', '')

            # Convert to list for manipulation
            preferred_locations = preferred_locations_str.split(',') if preferred_locations_str else []

            if new_location in preferred_locations:
                return jsonify({"error": "Location already added"}), 400

            # Add new location
            preferred_locations.append(new_location)

            # Convert back to comma-separated string
            updated_locations_str = ','.join(preferred_locations)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_locations': updated_locations_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Location added successfully", "locations": preferred_locations}), 200

        elif request.method == 'DELETE':
            data = request.get_json()
            location_to_remove = data.get('location')

            if not location_to_remove:
                return jsonify({"error": "Location is required for deletion"}), 400

            # Fetch current preferred locations
            response = supabase.table('user_job_preferences').select('preferred_locations').eq('user_id', current_user.id).single().execute()
            preferred_locations_str = response.data.get('preferred_locations', '')

            # Convert to list for manipulation
            preferred_locations = preferred_locations_str.split(',') if preferred_locations_str else []

            if location_to_remove not in preferred_locations:
                return jsonify({"error": "Location not found"}), 400

            # Remove location
            preferred_locations.remove(location_to_remove)

            # Convert back to comma-separated string
            updated_locations_str = ','.join(preferred_locations)

            # Update database
            supabase.table('user_job_preferences').update({
                'preferred_locations': updated_locations_str
            }).eq('user_id', current_user.id).execute()

            return jsonify({"message": "Location removed successfully", "locations": preferred_locations}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@main_bp.route('/update-relocation-preference', methods=['POST'])
@login_required
def update_relocation_preference():
    data = request.get_json()
    willing_to_relocate = data.get('willing_to_relocate')

    if not willing_to_relocate:
        return jsonify({"error": "Relocation preference is required"}), 400

    try:
        # Update the relocation preference in the database
        response = supabase.table('user_job_preferences').update({
            'willing_to_relocate': willing_to_relocate
        }).eq('user_id', current_user.id).execute()

        if response.data:
            return jsonify({"message": "Relocation preference updated successfully"}), 200
        else:
            return jsonify({"error": "Failed to update relocation preference"}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# @main_bp.route('/job_preferences', methods=['GET', 'POST'])
# @login_required
# def job_preferences():
#     form = JobPreferencesForm()
#     education_form = EducationEntryForm()
#     profile_id = current_user.id
#     is_subscribed = current_user.is_subscribed if not current_user.is_anonymous else False
#     job_preferences_response = supabase.table('user_job_preferences').select('*').eq('user_id', profile_id).execute()
#     job_preferences_data = job_preferences_response.data[0] if job_preferences_response.data else {}
#     form.full_name.data = job_preferences_data.get('real_name', '')
#     form.ideal_work_situation.data = job_preferences_data.get('ideal_work_situation', '')
#     form.preferred_industries.data = job_preferences_data.get('preferred_industries', '')
#     form.preferred_roles_responsibilities.data = job_preferences_data.get('preferred_roles_responsibilities', '')
#     form.work_arrangement_preference.data = job_preferences_data.get('work_arrangement_preference', '')
#     form.current_state.data = job_preferences_data.get('current_state', '')
#     form.current_city.data = job_preferences_data.get('current_city', '')
#     form.current_address.data = job_preferences_data.get('current_address', '')
#     form.willing_to_relocate.data = job_preferences_data.get('willing_to_relocate', '')
#     relocation_preference_value = job_preferences_data.get('relocation_preference', None)
#     # If the value is None (NULL in the database), assign a default value
#     if relocation_preference_value is None:
#         relocation_preference_value = 'anywhere'  # or 'specific', depending on your default preference

#     # Ensure the value is one of the valid choices
#     if relocation_preference_value not in ['anywhere', 'specific']:
#         relocation_preference_value = 'anywhere'  # Fallback to 'anywhere' if the value is invalid

#     form.relocation_preference.data = relocation_preference_value
#     form.preferred_locations.data = job_preferences_data.get('preferred_locations', '')
#     form.expected_salary_range.data = job_preferences_data.get('expected_salary_range', '')
#     # form.industry_importance.data = job_preferences_data.get('industry_importance', 0)
#     # form.location_work_arrangement_importance.data = job_preferences_data.get('location_work_arrangement_importance', 0)
#     # form.role_responsibilities_importance.data = job_preferences_data.get('role_responsibilities_importance', 0)
#     # form.salary_importance.data = job_preferences_data.get('salary_importance', 0)
#     # form.company_prestige_importance.data = job_preferences_data.get('company_prestige_importance', 0)
#     form.postnomial.data = job_preferences_data.get('postnomial', '')
#     form.contact_phone.data = job_preferences_data.get('phone', '')
#     form.contact_email.data = job_preferences_data.get('email', '')
#     work_experience = supabase.table('work_experience').select('*').eq('profile_id', profile_id).execute()
#     education = supabase.table('education').select('*').eq('profile_id', profile_id).execute()
#     certifications = supabase.table('certifications').select('*').eq('profile_id', profile_id).execute()

#     if form.validate_on_submit():
#         print(f"Form data on submission: {form.data}")

#         form.preferred_industries.data = form.preferred_industries.data.replace("[", "").replace("]", "").replace('"',"")
#         form.preferred_roles_responsibilities.data = form.preferred_roles_responsibilities.data.replace("[", "").replace("]", "").replace('"',"")
#         # Parse the comma-separated strings into lists
#         preferred_industries = form.preferred_industries.data.split(',') if form.preferred_industries.data else None
#         preferred_roles= form.preferred_roles_responsibilities.data.split(',') if form.preferred_roles_responsibilities.data else None

#         # Optionally strip whitespace around the values
#         preferred_industries = [industry.strip() for industry in preferred_industries]
#         preferred_roles= [role.strip() for role in preferred_roles]


#         # # Handle living address fields
#         # current_country = form.current_country.data
#         current_state = form.current_state.data
#         current_city = form.current_city.data
#         current_address = form.current_address.data

#         if not (current_state and current_city and current_address):
#             flash("All fields under 'Living Address' are required.", "error")
#             return redirect(url_for('main.job_preferences'))

#         # Handle preferred locations
#         if form.willing_to_relocate.data.lower() == 'yes':
#             preferred_locations_data = form.preferred_locations.data
#             print(f"Preferred Locations Raw Data: '{preferred_locations_data}'")

#             if preferred_locations_data:
#                 try:
#                     preferred_locations = json.loads(preferred_locations_data)
#                     print(f"JSON Parsed Preferred Locations: {preferred_locations}")
#                 except json.JSONDecodeError:
#                     preferred_locations = [loc.strip() for loc in preferred_locations_data.split(',') if loc.strip()]
#                     print(f"Comma-Separated Parsed Preferred Locations: {preferred_locations}")
#             else:
#                 preferred_locations = ["United States"]
#                 flash("Assigned 'United States' as the default preferred location.", "info")
#         else:
#             preferred_locations = []

#         # Prepare the final data structure
#         # values = {
#         #     'user_id': current_user.id,
#         #     'ideal_work_situation': form.ideal_work_situation.data,
#         #     'preferred_industries': preferred_industries,
#         #     'preferred_roles_responsibilities': preferred_roles,
#         #     'work_arrangement_preference': form.work_arrangement_preference.data,
#         #     'current_state': current_state,
#         #     'current_city': current_city,
#         #     'current_address': current_address,
#         #     'willing_to_relocate': form.willing_to_relocate.data,
#         #     'relocation_preference': form.relocation_preference.data if form.willing_to_relocate.data.lower() == 'yes' else None,
#         #     'preferred_locations': preferred_locations,
#         #     'expected_salary_range': form.expected_salary_range.data,
#         #     'industry_importance': int(form.industry_importance.data),
#         #     'location_work_arrangement_importance': int(form.location_work_arrangement_importance.data),
#         #     'role_responsibilities_importance': int(form.role_responsibilities_importance.data),
#         #     'salary_importance': int(form.salary_importance.data),
#         #     'company_prestige_importance': int(form.company_prestige_importance.data),
#         # }

#         values = {
#             'user_id': current_user.id,
#             'real_name': form.full_name.data,
#             'ideal_work_situation': form.ideal_work_situation.data,
#             'preferred_industries': preferred_industries,
#             'preferred_roles_responsibilities': preferred_roles,
#             'work_arrangement_preference': form.work_arrangement_preference.data,
#             'current_state': current_state,
#             'current_city': current_city,
#             'current_address': current_address,
#             'willing_to_relocate': form.willing_to_relocate.data,
#             'relocation_preference': form.relocation_preference.data if form.willing_to_relocate.data.lower() == 'yes' else None,
#             'preferred_locations': preferred_locations,
#             'expected_salary_range': form.expected_salary_range.data,
#             # 'industry_importance': int(form.industry_importance.data),
#             # 'location_work_arrangement_importance': int(form.location_work_arrangement_importance.data),
#             # 'role_responsibilities_importance': int(form.role_responsibilities_importance.data),
#             # 'salary_importance': int(form.salary_importance.data),
#             # 'company_prestige_importance': int(form.company_prestige_importance.data),
#             'postnomial': form.postnomial.data,
#             'phone': form.contact_phone.data,
#             'email': form.contact_email.data,
#         }

#         print(f"Form data before insert/update: {values}")

#         try:
#             # Insert or update preferences in Supabase
#             response = supabase.table('user_job_preferences').select('*').eq('user_id', current_user.id).execute()

#             if response.data:
#                 # Update existing preferences
#                 update_response = supabase.table('user_job_preferences').update(values).eq('user_id', current_user.id).execute()


#                 preferences_id = response.data[0]['id']
#             else:
#                 # Insert new preferences
#                 insert_response = supabase.table('user_job_preferences').insert(values).execute()


#                 preferences_id = insert_response.data[0]['id']

#             # Enqueue the background task
#             process_job_preferences.delay(current_user.id)
#             logger.info(f"Enqueued job_preferences processing for user {current_user.id}")

#             # Provide immediate feedback
#             flash('Job preferences updated successfully! Processing in the background.', 'success')
#             return redirect(url_for('main.index'))

#         except Exception as e:
#             flash(f"Error saving preferences: {str(e)}", 'error')
#             logger.error(f"Error saving preferences for user {current_user.id}: {str(e)}")
#             return redirect(url_for('main.job_preferences'))




#     else:
#         for field_name, errors in form.errors.items():
#             for error in errors:
#                 flash(f"{field_name.capitalize()}: {error}", "error")




        

#     return render_template('job_preferences.html', form=form, education_form=education_form, work_experience_entries=work_experience.data, education_entries=education.data, certification_entries=certifications.data, is_subscribed=is_subscribed)

@main_bp.route('/preferred-suggestions/locations', methods=['GET'])
def location_suggestions():
    query = request.args.get('q', '')

    # Example static suggestions (this could be a database query)
    all_locations = ["New York, NY", "Los Angeles, CA", "Chicago, IL", "Houston, TX", "Phoenix, AZ",
    "Philadelphia, PA", "San Antonio, TX", "San Diego, CA", "Dallas, TX", "San Jose, CA",
    "Austin, TX", "Jacksonville, FL", "San Francisco, CA", "Columbus, OH", "Fort Worth, TX",
    "Indianapolis, IN", "Charlotte, NC", "Seattle, WA", "Denver, CO", "El Paso, TX",
    "Washington, DC", "Boston, MA", "Detroit, MI", "Nashville, TN", "Oklahoma City, OK",
    "Portland, OR", "Las Vegas, NV", "Louisville, KY", "Baltimore, MD", "Milwaukee, WI",
    "Albuquerque, NM", "Tucson, AZ", "Fresno, CA", "Sacramento, CA", "Long Beach, CA",
    "Kansas City, MO", "Mesa, AZ", "Virginia Beach, VA", "Atlanta, GA", "Colorado Springs, CO",
    "Omaha, NE", "Raleigh, NC", "Miami, FL", "Cleveland, OH", "Tulsa, OK",
    "Oakland, CA", "Minneapolis, MN", "Wichita, KS", "New Orleans, LA", "Arlington, TX",
    "Bakersfield, CA", "Tampa, FL", "Honolulu, HI", "Anaheim, CA", "Santa Ana, CA",
    "Corpus Christi, TX", "Riverside, CA", "St. Louis, MO", "Pittsburgh, PA", "Greensboro, NC",
    "Lincoln, NE", "Anchorage, AK", "Plano, TX", "Orlando, FL", "Irvine, CA",
    "Laredo, TX", "Chula Vista, CA", "Durham, NC", "Jersey City, NJ", "Fort Wayne, IN",
    "St. Petersburg, FL", "Chandler, AZ", "Lubbock, TX", "Madison, WI", "Gilbert, AZ",
    "Reno, NV", "Hialeah, FL", "Baton Rouge, LA", "Richmond, VA", "Boise, ID",
    "San Bernardino, CA", "Spokane, WA", "Des Moines, IA", "Modesto, CA", "Fremont, CA",
    "Santa Clarita, CA", "Mobile, AL", "Oxnard, CA", "Moreno Valley, CA", "Huntington Beach, CA",
    "Aurora, CO", "Columbia, SC", "Grand Rapids, MI", "Salt Lake City, UT", "Tallahassee, FL",
    "Overland Park, KS", "Knoxville, TN", "Worcester, MA", "Newport News, VA", "Brownsville, TX",
    "Santa Rosa, CA", "Vancouver, WA", "Fort Lauderdale, FL", "Sioux Falls, SD", "Ontario, CA"]

    # Filter based on query
    suggestions = [loc for loc in all_locations if query.lower() in loc.lower()]

    return jsonify(suggestions), 200

@main_bp.route('/preferred-suggestions/industries', methods=['GET'])
def industry_suggestions():
    query = request.args.get('q', '')

    # Example static suggestions for industries
    all_industries = [
    # // Technology and IT
    "Software Development", "Cloud Computing", "Artificial Intelligence", "Machine Learning",
    "Cybersecurity", "Data Analytics", "Internet of Things (IoT)", "Blockchain",
    "Quantum Computing", "Robotics", "Virtual Reality", "Augmented Reality",
    "5G Technology", "Edge Computing", "DevOps", "IT Consulting",

    # // Healthcare and Life Sciences
    "Pharmaceuticals", "Biotechnology", "Medical Devices", "Healthcare IT",
    "Telemedicine", "Genomics", "Personalized Medicine", "Mental Health Services",
    "Elder Care", "Veterinary Medicine", "Dental Care", "Physical Therapy",
    "Nutrition and Wellness", "Health Insurance", "Medical Research", "Public Health",

    # // Finance and Banking
    "Commercial Banking", "Investment Banking", "Asset Management", "Hedge Funds",
    "Venture Capital", "Private Equity", "Insurance", "Financial Technology (FinTech)",
    "Cryptocurrency", "Personal Finance", "Accounting", "Tax Services",
    "Real Estate Investment", "Mortgage Lending", "Credit Services", "Financial Consulting",

    # // Education and Training
    "K-12 Education", "Higher Education", "Online Education", "EdTech",
    "Corporate Training", "Language Learning", "Special Education", "Early Childhood Education",
    "Vocational Training", "Test Preparation", "Educational Consulting", "Tutoring Services",
    "Adult Education", "STEM Education", "Art Education", "Music Education",

    # // Manufacturing and Industry
    "Automotive Manufacturing", "Aerospace Manufacturing", "Electronics Manufacturing",
    "Textile Manufacturing", "Food Processing", "Chemical Manufacturing",
    "Pharmaceutical Manufacturing", "Industrial Automation", "3D Printing",
    "Packaging", "Metalworking", "Plastics Manufacturing", "Paper and Pulp",
    "Machinery Manufacturing", "Furniture Manufacturing", "Toy Manufacturing",

    # // Retail and E-commerce
    "Online Retail", "Brick-and-Mortar Retail", "Grocery", "Fashion and Apparel",
    "Luxury Goods", "Consumer Electronics", "Home Improvement", "Sporting Goods",
    "Beauty and Cosmetics", "Pet Supplies", "Jewelry", "Books and Media",
    "Furniture and Home Decor", "Office Supplies", "Automotive Retail", "Specialty Foods",

    # // Hospitality and Tourism
    "Hotels and Resorts", "Restaurants", "Fast Food", "Catering",
    "Travel Agencies", "Airlines", "Cruise Lines", "Theme Parks",
    "Casinos and Gaming", "Event Planning", "Tour Operators", "Vacation Rentals",
    "Spa and Wellness Centers", "Timeshare", "Eco-Tourism", "Cultural Tourism",

    # // Energy and Utilities
    "Oil and Gas", "Renewable Energy", "Solar Power", "Wind Power",
    "Hydroelectric Power", "Nuclear Energy", "Energy Storage", "Smart Grid Technology",
    "Waste Management", "Water Treatment", "Natural Gas Distribution", "Electric Utilities",
    "Energy Efficiency", "Geothermal Energy", "Biomass Energy", "Hydrogen Fuel Cells",

    # // Telecommunications
    "Wireless Carriers", "Broadband Providers", "Satellite Communications", "Fiber Optics",
    "Telecom Equipment", "VoIP Services", "Network Infrastructure", "Telecom Software",
    "Mobile Virtual Network Operators", "Unified Communications", "Telecom Consulting", "Data Centers",

    # // Automotive
    "Car Manufacturing", "Electric Vehicles", "Autonomous Vehicles", "Auto Parts Manufacturing",
    "Car Dealerships", "Auto Repair and Maintenance", "Fleet Management", "Car Rental",
    "Automotive Design", "Motorcycle Manufacturing", "Truck Manufacturing", "Automotive Software",

    # // Aerospace and Defense
    "Aircraft Manufacturing", "Space Technology", "Satellite Systems", "Defense Contracting",
    "Missile Systems", "Military Vehicles", "Avionics", "Drone Technology",
    "Air Traffic Control Systems", "Aircraft Maintenance", "Space Exploration", "Rocket Propulsion",

    # // Agriculture and Farming
    "Crop Farming", "Livestock Farming", "Organic Farming", "Precision Agriculture",
    "Aquaculture", "Forestry", "Agricultural Biotechnology", "Farm Equipment",
    "Seed Technology", "Pesticides and Fertilizers", "Vertical Farming", "Hydroponics",
    "Agricultural Drones", "Food Safety", "Sustainable Agriculture", "Agritourism",

    # // Construction and Real Estate
    "Residential Construction", "Commercial Construction", "Infrastructure Development",
    "Architecture", "Civil Engineering", "Real Estate Development", "Property Management",
    "Interior Design", "Landscape Architecture", "Building Materials", "Smart Home Technology",
    "Green Building", "Facilities Management", "Urban Planning", "Real Estate Investment Trusts",

    # // Entertainment and Media
    "Film Production", "Television Broadcasting", "Streaming Services", "Music Industry",
    "Video Game Development", "Publishing", "Advertising", "Public Relations",
    "Social Media", "News Media", "Animation", "Podcasting",
    "Live Events", "Sports Entertainment", "Radio Broadcasting", "Digital Marketing",

    # // Transportation and Logistics
    "Trucking", "Rail Transport", "Air Cargo", "Maritime Shipping",
    "Logistics Software", "Warehousing", "Supply Chain Management", "Last-Mile Delivery",
    "Freight Forwarding", "Autonomous Logistics", "Cold Chain Logistics", "Reverse Logistics",
    "Intermodal Transportation", "Logistics Consulting", "Postal Services", "Courier Services",

    # // Environmental Services
    "Environmental Consulting", "Pollution Control", "Recycling", "Green Technology",
    "Climate Change Mitigation", "Conservation", "Sustainable Development", "Ecological Restoration",
    "Environmental Impact Assessment", "Hazardous Waste Management", "Air Quality Management", "Water Conservation",

    # // Professional Services
    "Legal Services", "Management Consulting", "Human Resources", "Recruitment",
    "Marketing Services", "Graphic Design", "Engineering Services", "Market Research",
    "Business Process Outsourcing", "Translation Services", "Data Entry Services", "Transcription Services",

    # // Non-Profit and Social Services
    "Charitable Organizations", "Social Advocacy", "Community Development", "Disaster Relief",
    "International Aid", "Human Rights", "Animal Welfare", "Environmental Conservation",
    "Arts and Culture", "Education Foundations", "Healthcare Foundations", "Religious Organizations"]

    # Filter based on query
    suggestions = [industry for industry in all_industries if query.lower() in industry.lower()]

    return jsonify(suggestions), 200

@main_bp.route('/preferred-suggestions/roles', methods=['GET'])
def role_suggestions():
    query = request.args.get('q', '')

    # Example static suggestions for job roles
    all_roles = [
    # // Management Roles
    "Chief Executive Officer (CEO)", "Chief Financial Officer (CFO)", "Chief Operating Officer (COO)",
    "Chief Technology Officer (CTO)", "Chief Marketing Officer (CMO)", "Chief Human Resources Officer (CHRO)",
    "Director of Operations", "Project Manager", "Program Manager", "Department Manager",
    "Team Leader", "Supervisor",

    # // Finance and Accounting
    "Financial Analyst", "Accountant", "Auditor", "Budget Analyst", "Tax Specialist",
    "Investment Banker", "Financial Planner", "Risk Manager", "Actuary",

    # // Technology and IT
    "Software Engineer", "Full Stack Developer", "Front-end Developer", "Back-end Developer",
    "Data Scientist", "Database Administrator", "Systems Administrator", "Network Engineer",
    "Cloud Architect", "DevOps Engineer", "Information Security Analyst", "UI/UX Designer",
    "Machine Learning Engineer", "Artificial Intelligence Specialist", "QA Engineer",

    # // Sales and Marketing
    "Sales Representative", "Account Manager", "Business Development Manager",
    "Marketing Specialist", "Digital Marketing Manager", "Content Marketing Manager",
    "Brand Manager", "Product Marketing Manager", "SEO Specialist", "Social Media Manager",
    "Public Relations Specialist", "Market Research Analyst",

    # // Human Resources
    "HR Manager", "Recruiter", "Training and Development Specialist",
    "Compensation and Benefits Analyst", "Employee Relations Specialist",
    "Talent Acquisition Manager", "HR Business Partner", "Diversity and Inclusion Specialist",

    # // Operations and Logistics
    "Operations Manager", "Supply Chain Manager", "Logistics Coordinator",
    "Procurement Specialist", "Inventory Manager", "Quality Assurance Manager",
    "Facilities Manager", "Production Planner", "Process Improvement Specialist",

    # // Customer Service
    "Customer Service Representative", "Customer Success Manager",
    "Technical Support Specialist", "Client Relations Manager",

    # // Research and Development
    "Research Scientist", "Product Developer", "R&D Manager",
    "Innovation Specialist", "Patent Specialist",

    # // Legal
    "Corporate Lawyer", "Legal Counsel", "Compliance Officer",
    "Paralegal", "Intellectual Property Specialist",

    # // Healthcare
    "Physician", "Nurse", "Pharmacist", "Medical Researcher",
    "Healthcare Administrator", "Physical Therapist", "Occupational Therapist",

    # // Education
    "Teacher", "Professor", "Education Administrator", "Curriculum Developer",
    "Instructional Designer", "School Counselor", "Special Education Specialist",

    # // Creative and Design
    "Graphic Designer", "Art Director", "Copywriter", "Video Editor",
    "UX/UI Designer", "Product Designer", "Industrial Designer",

    # // Engineering
    "Mechanical Engineer", "Civil Engineer", "Electrical Engineer",
    "Chemical Engineer", "Aerospace Engineer", "Environmental Engineer",

    # // Consulting
    "Management Consultant", "Strategy Consultant", "IT Consultant",
    "Financial Consultant", "Human Resources Consultant",

    # // Data and Analytics
    "Data Analyst", "Business Intelligence Analyst", "Data Engineer",
    "Statistician", "Operations Research Analyst",

    # // Project Management
    "Project Coordinator", "Scrum Master", "Agile Coach",
    "Program Director", "Portfolio Manager",

    # // Communications
    "Communications Specialist", "Technical Writer", "Translator",
    "Interpreter", "Copyeditor",

    # // Specialized Roles
    "Sustainability Officer", "Diversity and Inclusion Manager",
    "Change Management Specialist", "Crisis Management Specialist",
    "Innovation Manager", "Knowledge Management Specialist"]

    # Filter based on query
    suggestions = [role for role in all_roles if query.lower() in role.lower()]

    return jsonify(suggestions), 200




@main_bp.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    profile_id = current_user.id

    if request.method == 'GET':
        # Fetch profile data
        profile_response = supabase.table('user_job_preferences').select('*').eq('user_id', current_user.id).execute()
        profile_data = profile_response.data[0] if profile_response.data else {}
        return jsonify(profile_data)

    if request.method == 'POST':
        data = request.json
        required_fields = ["first_name", "last_name", "profile_name", "postnomial", "email", "phone", "current_city", "current_state"]
        
        # Check for empty required fields
        errors = {field: "This field is required." for field in required_fields if not data.get(field, "").strip()}
        if errors:
            return jsonify({"error": "Validation failed", "fields": errors}), 400  # Return error with missing fields

        values = {
            'user_id': current_user.id,
            'real_name': f"{data['first_name'].strip()} {data['last_name'].strip()}".strip(),
            'profile_name': data['profile_name'].strip(),
            'postnomial': data['postnomial'].strip(),
            'email': data['email'].strip(),
            'phone': data['phone'].strip(),
            'current_city': data['current_city'].strip(),
            'current_state': data['current_state'].strip(),
        }

        try:
            response = supabase.table('user_job_preferences').select('*').eq('user_id', current_user.id).execute()
            if response.data:
                supabase.table('user_job_preferences').update(values).eq('user_id', current_user.id).execute()
                update_response = supabase.table('profiles').update({
                    'is_first_login': False
                }).eq('id', profile_id).execute()
                session['show_profile_modal'] = False
            else:
                supabase.table('user_job_preferences').insert(values).execute()
                update_response = supabase.table('profiles').update({
                    'is_first_login': False
                }).eq('id', profile_id).execute()
                session['show_profile_modal'] = False


            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500



@main_bp.route('/job_filters', methods=['GET', 'POST'])
@login_required
def job_filters():
    form = JobFiltersForm()
    profile_id = current_user.id
    job_filters_response = supabase.table('user_job_preferences').select('*').eq('user_id', current_user.id).execute()
    job_filters_data = job_filters_response.data[0] if job_filters_response.data else {}

    # Pre-fill form data
    form.ideal_work_situation.data = job_filters_data.get('ideal_work_situation', '')
    form.preferred_industries.data = job_filters_data.get('preferred_industries', '')
    form.work_arrangement_preference.data = job_filters_data.get('work_arrangement_preference', '')
    form.willing_to_relocate.data = job_filters_data.get('willing_to_relocate', '')
    form.relocation_preference.data = job_filters_data.get('relocation_preference', 'anywhere')
    form.preferred_locations.data = job_filters_data.get('preferred_locations', '')
    form.expected_salary_range.data = job_filters_data.get('expected_salary_range', '')

    if form.validate_on_submit():
        # Parse and clean form data
        preferred_industries = [
            industry.strip() for industry in form.preferred_industries.data.split(',') if industry.strip()
        ]
        preferred_roles = [
            role.strip() for role in form.preferred_roles_responsibilities.data.split(',') if role.strip()
        ]
        preferred_locations = json.loads(form.preferred_locations.data) \
            if form.preferred_locations.data else ["United States"]

        values = {
            'user_id': current_user.id,
            'ideal_work_situation': form.ideal_work_situation.data,
            'preferred_industries': preferred_industries,
            'preferred_roles_responsibilities': preferred_roles,
            'work_arrangement_preference': form.work_arrangement_preference.data,
            'willing_to_relocate': form.willing_to_relocate.data,
            'relocation_preference': form.relocation_preference.data if form.willing_to_relocate.data.lower() == 'yes' else None,
            'preferred_locations': preferred_locations,
            'expected_salary_range': form.expected_salary_range.data,
        }

        try:
            # Update or insert job filter data
            response = supabase.table('user_job_preferences').select('*').eq('user_id', current_user.id).execute()
            if response.data:
                supabase.table('user_job_preferences').update(values).eq('user_id', current_user.id).execute()
            else:
                supabase.table('user_job_preferences').insert(values).execute()

            flash('Job filters updated successfully!', 'success')
            return redirect(url_for('main.job_filters'))
        except Exception as e:
            flash(f"Error saving job filters: {str(e)}", 'error')

    return render_template('job_filters.html', form=form)

@main_bp.route('/work_experience', methods=['GET'])
def get_work_experience():
    entries = supabase.table('work_experience').select('*').eq('profile_id', current_user.id).execute()
    return jsonify(entries.data)

@main_bp.route('/education', methods=['GET'])
def get_education():
    entries = supabase.table('education').select('*').eq('profile_id', current_user.id).execute()
    return jsonify(entries.data)

@main_bp.route('/certifications', methods=['GET'])
def get_certifications():
    entries = supabase.table('certifications').select('*').eq('profile_id', current_user.id).execute()
    return jsonify(entries.data)

# Create Work Experience Entry
@main_bp.route('/work_experience', methods=['POST'])
def add_work_experience():
    data = request.json

    # Check if 'Currently Working' (Present) checkbox is selected
    end_month = data.get('end_month')
    end_year = data.get('end_year')

    # Check if 'Present' is selected
    if data.get('end_date') == 'Present':
        # If 'Present' is selected, set end_month and end_year to None or 0
        end_month = None
        end_year = None

    # Check the number of existing work experience entries for the current user
    existing_entries = supabase.table('work_experience').select('*').eq('profile_id', current_user.id).execute()

    # If the user already has 3 or more entries, show a flash message and redirect
    if len(existing_entries.data) >= 5:
        flash("You can only add up to 5 work experience entries.", "error")
        # Redirect to the work experience page or wherever you want to show the message
        return redirect(url_for('main.job_preferences'))  # Replace with the actual route name


    work_experience = {
        "profile_id": current_user.id,
        "company": data['company'],
        "title": data['title'],
        "description": data.get('description'),
        "start_month": data['start_month'],
        "start_year": data['start_year'],
        "end_month": data.get('end_month'),
        "end_year": data.get('end_year'),
    }
    response = supabase.table('work_experience').insert(work_experience).execute()
    return jsonify(response.data), 201

# Create Education Entry
@main_bp.route('/education', methods=['POST'])
def add_education():
    data = request.json

    # Check if 'Currently Studying' (Present) checkbox is selected
    end_month = data.get('end_month')
    end_year = data.get('end_year')

    # Check if 'Present' is selected
    if data.get('end_date') == 'Present':
        # If 'Present' is selected, set end_month and end_year to None or 0
        end_month = None
        end_year = None

    # Check the number of existing work experience entries for the current user
    existing_entries = supabase.table('education').select('*').eq('profile_id', current_user.id).execute()

    # If the user already has 3 or more entries, show a flash message and redirect
    if len(existing_entries.data) >= 3:
        flash("You can only add up to 3 education entries.", "error")
        # Redirect to the work experience page or wherever you want to show the message
        return redirect(url_for('main.job_preferences'))  # Replace with the actual route name


    education = {
        "profile_id": current_user.id,
        "institution": data['institution'],
        "degree": data['degree'],
        "degree_title": data['degree_title'],
        "field_of_study": data.get('field_of_study'),
        "start_month": data['start_month'],
        "start_year": data['start_year'],
        "end_month": data.get('end_month'),
        "end_year": data.get('end_year'),
    }
    response = supabase.table('education').insert(education).execute()
    return jsonify(response.data), 201

# Create Certification Entry
@main_bp.route('/certifications', methods=['POST'])
def add_certification():
    data = request.json

    # Check the number of existing work experience entries for the current user
    existing_entries = supabase.table('certifications').select('*').eq('profile_id', current_user.id).execute()

    # If the user already has 3 or more entries, show a flash message and redirect
    if len(existing_entries.data) >= 5:
        flash("You can only add up to 5 certification entries.", "error")
        # Redirect to the work experience page or wherever you want to show the message
        return redirect(url_for('main.job_preferences'))  # Replace with the actual route name

    certification = {
        "profile_id": current_user.id,
        "title": data['title'],
        "issuer": data.get('issuer'),
        "acquired_date": data['acquired_date'],
    }
    response = supabase.table('certifications').insert(certification).execute()
    return jsonify(response.data), 201

# Delete Work Experience Entry
@main_bp.route('/work_experience/<int:entry_id>', methods=['DELETE'])
def delete_work_experience(entry_id):
    response = supabase.table('work_experience').delete().eq('id', entry_id).execute()
    
    # Check if the response data is empty (meaning the deletion was successful)
    if response.data:
        return jsonify({"message": "Work experience entry deleted"}), 200
    
    # If no data is returned, that means no record was found with the given id
    return jsonify({"error": "Entry not found"}), 404

# Delete Education Entry
@main_bp.route('/education/<int:entry_id>', methods=['DELETE'])
def delete_education(entry_id):
    response = supabase.table('education').delete().eq('id', entry_id).execute()
    if response.data:
        return jsonify({"message": "Education entry deleted"}), 200
    return jsonify({"error": "Entry not found"}), 404

# Delete Certification Entry
@main_bp.route('/certifications/<int:entry_id>', methods=['DELETE'])
def delete_certification(entry_id):
    response = supabase.table('certifications').delete().eq('id', entry_id).execute()
    if response.data:
        return jsonify({"message": "Certification entry deleted"}), 200
    return jsonify({"error": "Entry not found"}), 404




@main_bp.route('/create_resume_and_cover_letter', methods=['GET', 'POST'])
@login_required
def create_resume_and_cover_letter():
    #Items needed for this workflow:
    #uuid
    #job data (company, title, location, salary, job description)
    try:
        id = current_user.id
        response = requests.get(f"https://cognibly-n8n-hmasbrgge7gkdtew.southcentralus-01.azurewebsites.net/webhook/b1d63dbd-fb55-4707-bc9a-324130d3ce55?uuid={id}")
        return ("Sent workflow request to N8N.")
    except Exception as e:
        return (f"Failed because {str(e)}")


def init_drive_service():
    SERVICE_ACCOUNT_FILE = './e33096a34c3dab3b772ffda2300fe292/devcreating-land-1717741719209-eedafe3bf7d0.json'
    SCOPES = ['https://www.googleapis.com/auth/drive']

    credentials = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    return build('drive', 'v3', credentials=credentials)


def capitalize_first_letter(text):
    # Capitalize only the first letter of the text, leave the rest unchanged
    if text:
        return text[0].upper() + text[1:]
    return text  # Return the original text if it's empty

def replace_education_placeholder(document, education_data):
    # Define degree acronyms and titles
    degree_titles = {
        'BA': 'Bachelor of Arts',
        'BS': 'Bachelor of Science',
        'BSc': 'Bachelor of Science',
        'BBA': 'Bachelor of Business Administration',
        'BFA': 'Bachelor of Fine Arts',
        'BMus': 'Bachelor of Music',
        'BEd': 'Bachelor of Education',
        'BEng': 'Bachelor of Engineering',
        'BSN': 'Bachelor of Science in Nursing',
        'MA': 'Master of Arts',
        'MS': 'Master of Science',
        'MSc': 'Master of Science',
        'MBA': 'Master of Business Administration',
        'MFA': 'Master of Fine Arts',
        'MEd': 'Master of Education',
        'MPH': 'Master of Public Health',
        'MSW': 'Master of Social Work',
        'MMus': 'Master of Music',
        'MEng': 'Master of Engineering',
        'PhD': 'Doctor of Philosophy',
        'EdD': 'Doctor of Education',
        'DBA': 'Doctor of Business Administration',
        'MD': 'Doctor of Medicine',
        'JD': 'Juris Doctor (Law)',
        'DDS': 'Doctor of Dental Surgery',
        'DVM': 'Doctor of Veterinary Medicine',
        'PsyD': 'Doctor of Psychology',
        'PostDoc': 'Post-Doctoral Fellow/Research',
        'PDF': 'Post-Doctoral Fellowship'
    }

    # Iterate through paragraphs in the document to find the $EDUCATION placeholder
    for para in document.paragraphs:
        if '$EDUCATION' in para.text:  # Match the $EDUCATION placeholder
            # Clear the paragraph and replace with formatted education text
            for item in education_data:
                # Get degree abbreviation directly from item (e.g., 'MS')
                degree = item.get('degree_title', '')

                # Get the full degree title using the degree abbreviation
                degree_title = degree_titles.get(degree, '')

                # Capitalize the field of study and institution if available
                field_of_study = capitalize_first_letter(item['field_of_study']) if item.get('field_of_study') else ''
                institution = capitalize_first_letter(item['institution']) if item.get('institution') else ''

                # Format the text as "{degree} ({degree_title}) in {field_of_study} – {institution}"
                formatted_text = f"{degree} ({degree_title}) in {field_of_study} – {institution}" if field_of_study else f"{degree} ({degree_title}) – {institution}"

                # Clear the paragraph text (remove any existing content)
                para.clear()

                # Add the formatted text as a run with font size 10pt
                run = para.add_run(formatted_text)
                run.font.size = Pt(10)  # Set font size to 10pt

                # If you want to add additional lines after each entry, you could insert a new paragraph here
                # document.add_paragraph()  # Uncomment to add extra whitespace (new paragraph)

            break  # Exit the loop once the placeholder is replaced



def add_section_to_template(section_heading, data, key_map, document, is_certifications=False):
    # Iterate through paragraphs in the document to find the placeholder
    for para in document.paragraphs:
        if '$' + section_heading.upper() in para.text:  # Match the dynamic section heading placeholder
            # Replace the section heading with tab-stops formatted text
            insert_tab_stops_text(para, data, key_map, document, is_certifications)
            remove_paragraph(para)
            break  # Exit the loop once the placeholder is replaced

def remove_paragraph(para):
    # Access the underlying XML element of the paragraph
    para_element = para._element

    # Remove the paragraph from the document's XML structure
    para_element.getparent().remove(para_element)

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")  # Create a new XML element for the paragraph
    paragraph._p.addnext(new_p)  # Insert it after the current paragraph
    new_para = Paragraph(new_p, paragraph._parent)  # Wrap it as a `Paragraph` object
    if text:
        new_para.add_run(text)  # Add text if provided
    if style is not None:
        new_para.style = style  # Apply style if provided
    return new_para

def insert_tab_stops_text(para, data, key_map, document, is_certifications):
    """Insert tab stops and formatted text sequentially after the given paragraph."""
    for item in data:
        # Create a new paragraph for each entry
        new_para = insert_paragraph_after(para)

        left_text = item.get(key_map['left'], '')

        # Handle date formatting
        if is_certifications:
            date_text = item.get('acquired_date', 'Unknown Date')
        else:
            start_year = item.get('start_year', '')
            end_year = item.get('end_year', 'Present')
            if end_year == 'None' or end_year is None:
                end_year = 'Present'
            date_text = f"{start_year} - {end_year}"

        # Add left-aligned text with tab and right-aligned date
        run_left = new_para.add_run(left_text)
        run_left.font.size = Pt(11.5)

        run_tab_date = new_para.add_run(f"\t{date_text}")
        run_tab_date.font.size = Pt(10)

        # Calculate usable width for tab position
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin

        # Set the tab stop with right alignment
        new_para.paragraph_format.tab_stops.add_tab_stop(
            usable_width,
            alignment=WD_TAB_ALIGNMENT.RIGHT
        )

        # Handle additional lines (below_text)
        below_text = key_map.get('below', [])
        if isinstance(below_text, list):  # If below is a list of keys
            below_text_lines = [item.get(key, '') for key in below_text]
        else:
            below_text_lines = [item.get(below_text, '')]

        for below_line in below_text_lines:
            if below_line:  # Only add if there's text to add
                new_para.add_run(f"\n{below_line}").font.size = Pt(10)

        # Align the paragraph to the left (tab stop ensures the date aligns right)
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Manually insert a blank paragraph by adding a new empty paragraph after the current one
        new_blank_para = insert_paragraph_after(new_para)

        # Set a minimal formatting or style for the blank line if necessary
        new_blank_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_blank_para.add_run("")  # Just an empty run to create the blank line
        
def replace_placeholders(paragraphs, replacements):
    for paragraph in paragraphs:
        # Print the paragraph text for debugging
        print(f"Processing paragraph: {paragraph.text}")
        
        for placeholder, value in replacements.items():
            if value is None:
                value = ''  # Replace None with an empty string

            # Handle $FULLNAME special case
            if placeholder == '$FULLNAME':
                postnomial = replacements.get('$POSTNOMIAL', '')
                # Add a comma and postnomial if it's present
                replacement_text = value + (', ' + postnomial if postnomial else '')
                # Now replace the placeholder with the new value (full name + postnomial)
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement_text)
                        run.font.size = Pt(12)  # Set the font size for FULLNAME
            elif placeholder == '$POSTNOMIAL':
                # Regular replacement for $POSTNOMIAL
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        run.font.size = Pt(9)  # Set the font size for POSTNOMIAL
            elif placeholder == '$CONTACTDETAILS':
                # Regular replacement for $CONTACTDETAILS
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        run.font.size = Pt(11)  # Set the font size for CONTACTDETAILS
            else:
                # Regular placeholder replacement for all other placeholders
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        run.font.size = Pt(10)  # Default font size


def calculate_word_count(work_exp_count, education_count, certifications_count):
    # Work experience word count based on the inverted functionality
    static_edu_cert_word_count = (education_count + certifications_count) * 15

    work_exp_word_count = 0
    if work_exp_count > 5:
        work_exp_word_count = 96
    if work_exp_count == 5:
        work_exp_word_count = 108
    elif work_exp_count == 4:
        work_exp_word_count = 120
    elif work_exp_count == 3:
        work_exp_word_count = 160
    elif work_exp_count == 2:
        work_exp_word_count = 240
    elif work_exp_count == 1:
        work_exp_word_count = 480

    # Adjust work experience allocation based on remaining words after education and certifications
    work_exp_word_count = work_exp_word_count - static_edu_cert_word_count

    return work_exp_word_count


def generate_resume(job_data, user_id):
    # Fetch user data from Supabase
    user_data_response = supabase.table('user_job_preferences').select('*').eq('user_id', user_id).execute()
    user_data = user_data_response.data
    print(f"Job Data: {job_data}")

    if not user_data:
        flash("User data not found.", 'error')
        return redirect(url_for('some_error_page'))

    user_details = user_data[0]
    if not user_details.get('real_name'):
        flash("Real name is missing.", 'error')
        return redirect(url_for('some_error_page'))

    # Extract user details
    full_name = user_details['real_name']
    postnomial = user_details.get('postnomial', '')
    phone_number = user_details.get('phone', '')
    email = user_details.get('email', '')
    current_city = user_details.get('current_city', '')
    current_state = user_details.get('current_state', '')

    # Generate professional summary using OpenAI
    company = job_data['company']
    job_description = job_data.get('job_description', "")
    preferred_roles_responsibilities = user_details.get('preferred_roles_responsibilities', '')

    summary_prompt = f"""You are a professional resume summary writer. 
    The user will upload a job description they will use to apply for the job.

    Your task is to write a resume summary section that is no more than 60 words in length. 
    The summary must specifically be tailored for the job which the user provides as input. 

    Preferred Roles and Responsibilities: {preferred_roles_responsibilities}
    Applying to company: {company}
    """

    summary_response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.8,
        messages=[
            {"role": "system", "content": summary_prompt},
            {"role": "user", "content": f"Job Description: {job_description}"}
        ]
    )
    summary = summary_response.choices[0].message.content   

    skills_prompt = f"""The user will upload a job description they will use to apply for the job.

Your task is to write a list of skills no more than 45 words in length. The skills list must specifically be tailored for the job which the user provides as input. Use a mixture of hard skills (technologies) and soft skills, all of which should be related to the job.

The first letter of each skill must be Capitalized. Each skill must be separated by a comma and a space. List the skills in order of relevance to the job posting. 

IMPORTANT: Do not exceed 45 words in length. This is going at the bottom of a resume and if you exceed 45 words, it will not fit. Do not give the skills list a label like "Skills: " - Generate the skills list only, following the below Example List:
"""

    skills_response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.4,
        messages=[
            {"role": "system", "content": skills_prompt},
            {"role": "user", "content": f": {job_description}"}
        ]
    )
    skills = skills_response.choices[0].message.content  

    # Fetch additional details from Supabase
    work_experience = supabase.table('work_experience').select('*').eq('profile_id', user_id).execute()
    education = supabase.table('education').select('*').eq('profile_id', user_id).execute()
    certifications = supabase.table('certifications').select('*').eq('profile_id', user_id).execute()

    # Safely process data
    work_experience_data = [
        {
            "company": data.get('company', 'Unknown Company'),
            "title": data.get('title', 'Unknown Title'),
            "description": data.get('description', ''),
            "start_year": data.get('start_year', ''),
            "end_year": data.get('end_year', 'Present'),
        }
        for data in (work_experience.data or [])
    ]
    education_data = [
        {
            "institution": data.get('institution', 'Unknown Institution'),
            "degree": data.get('degree', 'Unknown Degree'),
            "degree_title": data.get('degree_title', 'Unknown Title'),
            "field_of_study": data.get('field_of_study', 'Unknown Field'),
            "start_year": data.get('start_year', 'Unknown'),
            "end_year": data.get('end_year', 'Present'),
        }
        for data in (education.data or [])
    ]
    certifications_data = [
        {
            "title": data.get('title', 'Unknown Certification'),
            "issuer": data.get('issuer', ''),
            "acquired_date": data.get('acquired_date', ''),
        }
        for data in (certifications.data or [])
    ]


 # For each work experience entry, decide whether to add more details based on the number of entries
    word_limit = calculate_word_count(len(work_experience_data), len(education_data), len(certifications_data))

    for work in work_experience_data:
        description_prompt = f"""You are a resume writer. Based on the title '{work['title']}' at '{work['company']}', write a description for this job role, as if you were the one that had it and want to give your objective outlook on it. Use maximum of {word_limit} words. Respect the word count. Do not give any headings or position headers.

Title: {work['title']}
Company: {work['company']}
"""

        description_response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.7,
            messages=[
                {"role": "system", "content": description_prompt},
                {"role": "user", "content": f"Job Description: {job_description}"}
            ]
        )

        # Extract the generated description
        generated_description = description_response.choices[0].message.content

        # Update the work entry with the adjusted description
        work['description'] = generated_description

    # Load the template
    try:
        template_path = os.path.join(os.path.dirname(__file__), 'docs', 'resume_template.docx')
        document = Document(template_path)
    except Exception as e:  
        flash("Template not found or invalid.", 'error')
        return redirect(url_for('some_error_page'))

    # Replace placeholders in the template
    placeholder_map = {
        "$FULLNAME": full_name,
        "$POSTNOMIAL": postnomial,
        "$CONTACTDETAILS": f"{current_city}, {current_state} · {email} · {phone_number}",
        "$SUMMARY": summary,
        "$SKILLS": skills,
    }

    # Replace placeholders in headers and footers
    for section in document.sections:
        # Process the header
        header = section.header
        replace_placeholders(header.paragraphs, placeholder_map)
        
        # Process the footer
        footer = section.footer
        replace_placeholders(footer.paragraphs, placeholder_map)

        replace_placeholders(document.paragraphs, placeholder_map)

    # Add sections
    add_section_to_template("WorkExperience", work_experience_data, {
        "left": "company",
        "below": ["title", "description"],
    }, document, False)

    replace_education_placeholder(document, education_data)

    add_section_to_template("Certifications", certifications_data, {
        "left": "issuer",
        "below": "title",
    }, document, True)

    # Save the document
    output_dir = os.path.join(os.getcwd(), 'static', 'cover_letter')
    os.makedirs(output_dir, exist_ok=True)

    job_title = job_data['job_title']
    output_filename = f"{job_title.replace(' ', '_')}_resume_{datetime.now().strftime('%I_%M_%d_%m_%Y')}.docx"
    output_path = os.path.join(output_dir, output_filename)

    # Save the DOCX document
    document.save(output_path)

    # Define PDF output path and filename
    pdf_output_filename = output_filename.replace('.docx', '.pdf')
    pdf_output_path = os.path.join(output_dir, pdf_output_filename)

    # Convert DOCX to PDF
    convert_docx_to_pdf(output_path, pdf_output_path)

    # Return the PDF directory and PDF filename
    return os.path.dirname(pdf_output_path), pdf_output_filename


def convert_docx_to_pdf(docx_path, pdf_path):
    # Make sure the LibreOffice binary is in your PATH
    # On Linux and macOS, the following command should work.
    # On Windows, use the full path to LibreOffice (e.g., 'C:/Program Files/LibreOffice/program/soffice.exe')
    
    libreoffice_command = [
        'libreoffice',  # Command to run LibreOffice
        '--headless',   # Run in headless mode (no GUI)
        '--convert-to', 'pdf',  # Convert to PDF
        '--outdir', os.path.dirname(pdf_path),  # Specify output directory
        docx_path  # Path to the input DOCX file
    ]
    
    # Run the command
    subprocess.run(libreoffice_command, check=True)
    print(f"PDF saved successfully at {pdf_path}")


def generate_cover_letter(job_data, user_id):
    # Fetch user data from Supabase
    user_data_response = supabase.table('user_job_preferences').select('*').eq('user_id', user_id).execute()
    user_data = user_data_response.data
    cover_letter_content = None

    if not user_data:
        flash("User data not found.", 'error')
        return redirect(url_for('some_error_page'))

    user_details = user_data[0]
    if user_details.get('real_name') is None:
        flash("Real name is missing", 'error')

    full_name = user_details['real_name']
    postnomial = user_details['postnomial']
    phone_number = user_details.get('phone', '')
    email = user_details['email']
    current_city = user_details['current_city']
    current_state = user_details['current_state']

    # Format contact details
    contact_details = f"{current_city}, {current_state}"
    if phone_number:
        contact_details += f" • {phone_number}"
    contact_details += f" • {email}"

    job_title = job_data['job_title']
    company = job_data['company']
    # Assuming `job_data` is available in the current context and contains the job description
    job_description = job_data.get('job_description', "")

    work_experience = supabase.table('work_experience').select('*').eq('profile_id', current_user.id).execute()
    education = supabase.table('education').select('*').eq('profile_id', current_user.id).execute()
    certifications = supabase.table('certifications').select('*').eq('profile_id', current_user.id).execute()
    
    # Extract relevant data from the responses
    work_experience_data = [
        {
            "company": data['company'],
            "title": data['title'],
            "description": data.get('description'),
            "start_month": data['start_month'],
            "start_year": data['start_year'],
            "end_month": data.get('end_month'),
            "end_year": data.get('end_year'),
        }
        for data in work_experience.data or []
    ]
    

    education_data = [
        {
            "institution": data['institution'],
            "degree": data['degree'],
            "field_of_study": data.get('field_of_study'),
            "start_month": data['start_month'],
            "start_year": data['start_year'],
            "end_month": data.get('end_month'),
            "end_year": data.get('end_year'),
        }
        for data in education.data or []
    ]

    certifications_data = [
        {
            "title": data['title'],
            "issuer": data.get('issuer'),
            "acquired_date": data['acquired_date'],
        }
        for data in certifications.data or []
    ]
    print(f"Job Description: f{job_description}, Company Name: {company}")
    summary_prompt = f"""You are a job application cover letter writer. The user will send a job description and company name and the resume they will use to apply for the job. Generate a cover letter of approximately 300 words in length which appropriately draws upon the experiences described in the user's resume to position the user as an excellent candidate for the job. 

Your writing style should draw from the Harvard Cover Letter Example I'm going to send below

IMPORTANT: Do not include any variables or fields which may require user input. Do not include a date. Do not include the recipient address. Instead of starting with "Dear [Recipient Name]" you must write something like "Dear Hiring Committee"

Sign the cover letter using details extrapolated from the user's resume. If you are unable to determine the details, you must use generic information. 

Please output the cover letter responses without annotations, footnotes, or bracketed comments. Generate the letter only. Do not provide any intro or summary after generating the letter.

Cover Letter Example:
Dear Morgan Smith:
I am a senior at Harvard College studying History and Literature. I am writing to apply for the
Marketing and Communications position at Jumpstart posted in Harvard’s Crimson Careers
database. I'm very excited about the field of education, and would welcome the opportunity to
bring my strong communication skills, creativity, and marketing experience to your growing
team.
Jumpstart's commitment to early education for every child is of particular interest to me because
of my passion for youth development. This past summer, I worked as a senior counselor in the
Summer Urban Program, which is dedicated to preventing summer learning loss for children in
the Boston and Cambridge area. I designed and taught fun, interactive classes to a group of 10
fifth graders, and planned and led local field trips and workshops daily with a junior counselor.
Throughout the summer, I consistently strived to create math, science, and reading lessons and
activities that were engaging and tailored to my students' needs.
Additionally, in my role as the Director of Marketing for the Social Innovation Collaborative, I
led our team in creating a social media strategy to drive our member recruitment efforts and
promote our programs and events on platforms including Facebook, Twitter, and Instagram.
With so many competing events on campus each day, I had to continually be creative in my
approach to developing and delivering content that would be compelling and effective. As a
result of my efforts, our group experienced a 20% increase in our membership base and a 15%
increase in our social media engagement. I’m excited at the prospect of bringing the skills I
developed through this experience to the Marketing and Communications role at Jumpstart.
Thank you for your consideration. I very much look forward to the opportunity to speak with you
in person about my interest in this position.

Work Experience: {work_experience}
Education: {education}
Certifications: {certifications}
Applying to company: {company}

Don't end the text with any "sincerely", "kind regards" or give any personal information at the end.
"""

    response = client.chat.completions.create(
    model="gpt-4o-mini", # model = "deployment_name".
    messages=[
        {"role": "system", "content": summary_prompt},
        {"role": "user", "content": f"Job Descriptions: {job_description}"}
    ]
    )

    cover_letter_content = response.choices[0].message.content   
    # # Generate company address
    # openai_url = f"{os.getenv('AZURE_OPENAI_ENDPOINT_COVER_LETTER')}/openai/deployments/gpt-4o-mini/chat/completions?api-version=2024-05-01-preview"
    # headers = {
    #     "Authorization": f"Bearer {token_provider.get_token()}",
    #     "Content-Type": "application/json"
    # }

    return_address_prompt = f"""
Provide the formatted U.S. address for the following company for use in a cover letter. 
The address should be in the following format:

Company Name  
Street Address  
City, State ZIP Code

If you do not know the address for this company, find it. If there are too many companies with the same name, just give the output only the name of the company and nothing else in this format:

Company Name
    """

    response = client.chat.completions.create(
    model="gpt-4o-mini", # model = "deployment_name".
    messages=[
        {"role": "system", "content": return_address_prompt},
        {"role": "user", "content": f"Company Name: {company}"}
    ]
    )

    company_address = response.choices[0].message.content


    # openai_payload_address = {
    #     "prompt": return_address_prompt,
    #     "max_tokens": 100,
    #     "temperature": 0.2
    # }
    # response_address = requests.post(openai_url, headers=headers, json=openai_payload_address)
    # if response_address.status_code != 200:
    #     flash("Error generating company address: " + response_address.text, 'error')
    #     return redirect(url_for('home'))

    # address_response = response_address.json()
    # company_address = address_response['choices'][0]['text'].strip()

    # Prepare placeholders for template replacement
    current_time = datetime.now().strftime("%m-%d-%Y")
    user_details_dict = {
        '$FULLNAME': full_name,
        '$POSTNOMIAL': postnomial,
        '$CONTACTDETAILS': contact_details,
        '$DATETIME': current_time,
        '$COMPANY_LOCATION': company_address,
        '$USER_LOCATION': f"{current_city}, {current_state}",
        '$SUMMARY': cover_letter_content,
        '$EMAIL': email,
        '$PHONENUMBER': phone_number
    }

    # Load and modify the Word template
    template_path = os.path.join(os.path.dirname(__file__), 'docs', 'cover_letter_template.docx')
    document = Document(template_path)
    for paragraph in document.paragraphs:
        for placeholder, value in user_details_dict.items():
            # Ensure that the value is a string, replacing None with an empty string
            if value is None:
                value = ''  # Or use a default string like 'N/A' if preferred
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
            
            # Now, update the font size for the paragraph's runs to 10Pt
            for run in paragraph.runs:
                run.font.size = Pt(10)


    # Replace placeholders in headers
    for section in document.sections:
        header = section.header
        replace_placeholders(header.paragraphs, user_details_dict)

    output_dir = os.path.join(os.getcwd(), 'static', 'cover_letter')
    os.makedirs(output_dir, exist_ok=True)

    output_filename = f"{job_title.replace(' ', '_')}_cover_letter_{datetime.now().strftime('%I_%M_%d_%m_%Y')}.docx"
    output_path = os.path.join(output_dir, output_filename)


    # Save the DOCX document
    document.save(output_path)

    # Define PDF output path and filename
    pdf_output_filename = output_filename.replace('.docx', '.pdf')
    pdf_output_path = os.path.join(output_dir, pdf_output_filename)

    # Convert DOCX to PDF
    convert_docx_to_pdf(output_path, pdf_output_path)

    # Return the PDF directory and PDF filename
    return os.path.dirname(pdf_output_path), pdf_output_filename



@main_bp.route('/generate_doc/<job_id>/<doc_type>', methods=['GET'])
@login_required
def generate_doc(job_id, doc_type):
    is_subscribed = current_user.is_subscribed if not current_user.is_anonymous else False

    if not is_subscribed:
        flash("You need to buy a subscription to use this feature.", "error")
        return redirect(url_for('main.jobs'))

    if not job_id or doc_type not in ['resume', 'cover_letter']:
        flash("Invalid parameters for document generation.", "error")
        return redirect(url_for('main.jobs'))

    try:
        # Fetch job details from the database using the job_id
        job_response = supabase.table('job_postings').select('*').eq('id', job_id).execute()
        job_data = job_response.data[0] if job_response.data else None
        print(job_data)

        if not job_data:
            raise ValueError("Job not found.")

        # Fetch the resume_template_path from the profiles table for the logged-in user
        user_id = current_user.id  # Assuming current_user.id gives the logged-in user's ID


        # Generate the document based on the type
        if doc_type == 'resume':
            document_dir, document_filename = generate_resume(job_data, user_id)
        elif doc_type == 'cover_letter':
            document_dir, document_filename = generate_cover_letter(job_data, user_id)

        print(document_dir, document_filename)
        # Serve the file for download
        return send_from_directory(
            directory=document_dir,
            path=document_filename,
            as_attachment=True
        )

    except Exception as e:
        flash("An error occurred while generating the document.", "error")
        logger.exception(f"Error generating document for job_id={job_id}: {str(e)}")
        return redirect(url_for('main.jobs'))

